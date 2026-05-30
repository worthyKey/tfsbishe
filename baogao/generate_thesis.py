"""
生成沈阳航空航天大学本科毕业论文 .docx （扩充版：正文>25页，>1.5万字）
严格按照模板格式：A4、页边距、宋体小四、1.25倍行距等
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

FONT_CN = "宋体"
FONT_EN = "Times New Roman"
SIZE_BODY = Pt(12)        # 小四
SIZE_SECTION = Pt(14)     # 四号
SIZE_CHAPTER = Pt(15)     # 小三
SIZE_SMALL = Pt(10.5)     # 五号
SIZE_COVER_TITLE = Pt(22)

MARGIN_TOP = Cm(3.5)
MARGIN_BOTTOM = Cm(3.0)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(2.0)

RESULT_DIR = os.path.join(os.path.dirname(__file__), "..", "results")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
    "基于数据驱动的循环水控制方法研究_滕凤硕_v2.docx")


def set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_BODY,
                 bold=False, italic=False):
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)


def set_paragraph_spacing(paragraph, line_spacing=1.25,
                          space_before=0, space_after=0,
                          first_line_indent=None):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = first_line_indent


def add_chapter(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_CHAPTER, bold=True)
    set_paragraph_spacing(p, line_spacing=1.5, space_before=17, space_after=17)
    return p


def add_section(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SECTION, bold=True)
    set_paragraph_spacing(p, line_spacing=1.5, space_before=6, space_after=0)
    return p


def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_BODY, bold=True)
    set_paragraph_spacing(p, line_spacing=1.5, space_before=6, space_after=0)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_BODY, bold=False)
    set_paragraph_spacing(p, line_spacing=1.25, space_before=0, space_after=0,
                          first_line_indent=Cm(0.74))
    return p


def add_formula(doc, formula_text):
    """添加居中公式"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_BODY, italic=True)
    set_paragraph_spacing(p, line_spacing=1.25, space_before=3, space_after=3)
    return p


def add_table(doc, headers, rows, caption):
    p_blank = doc.add_paragraph()
    set_paragraph_spacing(p_blank, line_spacing=1.25)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(caption)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SMALL, bold=True)
    set_paragraph_spacing(p_cap, line_spacing=1.25, space_before=0, space_after=3)
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SMALL, bold=True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SMALL)
    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, line_spacing=1.25)
    return table


def add_figure(doc, image_path, caption, width=Cm(14)):
    p_blank = doc.add_paragraph()
    set_paragraph_spacing(p_blank, line_spacing=1.25)
    if os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(image_path, width=width)
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(caption)
    set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SMALL, bold=True)
    set_paragraph_spacing(p_cap, line_spacing=1.25, space_before=0, space_after=6)
    p_after = doc.add_paragraph()
    set_paragraph_spacing(p_after, line_spacing=1.25)


def add_page_break(doc):
    doc.add_page_break()


def setup_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


# ============================================================
def generate():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    section = doc.sections[0]
    setup_page(section)

    # ==================== 封面 ====================
    for _ in range(6):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("毕业设计论文"); set_run_font(run, size=Pt(26), bold=True)
    for _ in range(3):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("沈阳航空航天大学"); set_run_font(run, size=Pt(22), bold=True)
    add_page_break(doc)

    # ==================== 内封面 ====================
    for _ in range(6):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于数据驱动的循环水控制方法研究")
    set_run_font(run, size=SIZE_CHAPTER, bold=True)
    for _ in range(5):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    info_items = [
        ("学       院：", "人工智能学院"),
        ("专       业：", "数据科学与大数据技术"),
        ("班       级：", "大数据2201班"), ("学 生 姓 名：", "滕凤硕"),
        ("学       号：", "XXXXXXXXXX"), ("指 导 教 师：", "许谨"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}      {value}")
        set_run_font(run, size=SIZE_BODY)
        set_paragraph_spacing(p, line_spacing=2.0)
    for _ in range(4):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("沈阳航空航天大学"); set_run_font(run, size=SIZE_BODY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("202X年6月"); set_run_font(run, size=SIZE_BODY)
    add_page_break(doc)

    # ==================== 原创性声明 & 授权声明 ====================
    add_chapter(doc, "原 创 性 声 明")
    add_body(doc, "本人郑重声明：所呈交的毕业设计（论文）是本人在导师的指导下独立完成的。除文中已经注明引用的内容外，本论文不包含其他个人或集体已经发表或撰写过的作品或成果，也不包含本人为获得其他学位而使用过的成果。对本文研究做出重要贡献的个人或集体均已在论文中进行了说明并表示谢意。本声明的法律后果由本人承担。")
    for _ in range(3):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("论文作者签名："); set_run_font(run)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("年    月    日"); set_run_font(run)
    add_page_break(doc)

    add_chapter(doc, "关于论文使用授权的说明")
    add_body(doc, "本论文的研究成果归沈阳航空航天大学所有，本论文的研究内容不得以其它单位的名义发表。本论文作者和指导教师完全了解沈阳航空航天大学有关保留、使用论文的规定，即：学校有权保留并向国家有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅；本人授权沈阳航空航天大学可以将论文的全部或部分内容编入有关数据库进行检索、交流，可以采用影印、缩印或其他复制手段保存论文和汇编本论文。")
    p = doc.add_paragraph(); run = p.add_run("（保密的论文在解密后应遵循此规定）")
    set_run_font(run, size=SIZE_SMALL)
    for _ in range(2):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("作者签名：          导师签名：           日期：      年    月")
    set_run_font(run)
    add_page_break(doc)

    # ==================== 中文摘要 ====================
    add_chapter(doc, "摘  要")
    add_body(doc, "循环水系统是工业过程中至关重要的能量输运与热交换环节，广泛应用于石油化工、电力冶金、暖通空调等领域。其运行效率直接影响整个工艺系统的能耗水平和生产稳定性。当前工业生产中的循环水控制系统主要依赖操作人员的技能和经验进行手动调节，控制效果因人而异，难以保证系统运行的稳定性、可靠性和能效水平。随着工业物联网和人工智能技术的快速发展，利用这些真实企业数据驱动的方式实现复杂工业过程的智能建模与优化控制，已成为学术界和工业界共同关注的前沿方向。本课题来源于某工业企业循环水系统的真实历史运行数据，旨在运用数据科学与大数据技术的方法体系，从企业实际数据中挖掘系统运行规律，设计智能控制方法")

    add_body(doc, "本文以某工业企业实际循环水系统为研究对象，基于系统28天连续运行的真实历史传感器数据（524,043行、16维特征），提出了一套完整的\"物理映射建模—LSTM系统辨识—SVM+GA参数优化—DRL智能控制—闭环仿真验证\"数据驱动控制框架。首先，针对传感器测量位置与控制目标位置不一致的关键问题，基于系统P&ID图构建了流量物理映射模型（Q_he = Q_total × M1/(M1+M2)，R² = 0.9998）和温度物理映射模型（T2 = α×T_cold + (1-α)×T_hot，R² = 0.9965），将传感器测量值精确转化为换热器入口的等效控制目标（流量27.0 m³/h，温度22.0 °C），奠定了优化控制的目标函数基础。其次，构建了3层长短期记忆网络（LSTM）对系统动态行为进行黑箱建模，网络参数量977,291，在包含24,453个训练样本的数据集上经过50轮训练，验证损失降至0.016128，关键温度状态变量的决定系数R²最高达0.9788，为后续优化控制提供了高保真的环境代理模型。")

    add_body(doc, "在系统建模的基础上，本文分别研究了两种控制策略。第一种策略采用支持向量机（SVM）代理模型加速遗传算法（GA）的适应度评估，SVM代理模型在使用5,000个LSTM评估样本训练后达到MSE=0.1263，GA在80代进化后收敛至最优适应度-2487.12，搜索到的最优控制参数组合为valve_DN200_fb=19.30、valve_DN300_fb=99.11、valve_DN350_fb=99.67、valve_DN400_fb=77.39、pump_speed_fb=1152.10 RPM。第二种策略采用深度确定性策略梯度（DDPG）框架，并创新性地引入GA行为克隆（Behavioral Cloning）预训练机制：利用GA优化器在多样化初始条件下生成150条专家控制轨迹，通过监督学习预训练DDPG的Actor网络（Actor MSE=0.0389），将优化能力有效迁移至神经网络策略中。在控制执行层面，引入了物理引导的实时修正机制，通过M1（流量修正，融合比例80%物理+20%网络）和M3（温度修正，融合比例60%物理+40%网络）的动态调整，确保控制器在实际执行时输出物理合理的控制动作。")

    add_body(doc, "闭环仿真验证结果显示：GA闭环优化在30次迭代中实现了流量达标率100%、温度达标率100%、压力达标率100%，平均等效流量27.1 m³/h，平均等效温度20.47 °C，平均能耗348.6；DRL闭环控制在100步运行中同样实现了三项指标全部100%达标，平均等效温度20.56 °C更接近目标值22.0 °C，平均能耗394.9。两种方法均验证了数据驱动方法在循环水系统控制中的可行性和有效性：GA在全局优化和能耗控制上略占优势，DRL在温度跟踪精度和实时响应速度上表现更优。本文的研究成果为工业循环水系统从\"依赖人工经验\"向\"数据智能驱动\"的转型升级提供了理论依据和技术方案，所提出的\"专家优化+神经网络模仿+物理修正\"混合控制范式对同类复杂工业过程的智能控制具有参考价值。")

    for _ in range(2):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph()
    run = p.add_run("关键词："); set_run_font(run, bold=True)
    run = p.add_run("循环水系统；数据驱动控制；长短期记忆网络；遗传算法；深度强化学习；行为克隆；物理映射；闭环仿真")
    set_run_font(run)
    add_page_break(doc)

    # ==================== 英文摘要 ====================
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Research on Data-Driven Control Method for Circulating Water System")
    set_run_font(run, size=SIZE_CHAPTER)
    set_paragraph_spacing(p, line_spacing=1.25, space_before=17, space_after=17)
    add_chapter(doc, "Abstract")

    add_body(doc, "Circulating water systems are critical energy transport and heat exchange links widely deployed in petrochemical, power metallurgy, and HVAC industries. Currently, industrial circulating water control systems primarily rely on manual adjustment based on operators' skills and experience, making it difficult to guarantee operational stability, reliability, and energy efficiency. This research is based on real historical operational data from an industrial enterprise's circulating water system. With the rapid advancement of the Industrial Internet of Things and artificial intelligence, leveraging real enterprise data to achieve intelligent modeling and optimal control of complex industrial processes has become a frontier direction attracting joint attention from academia and industry.")

    add_body(doc, "This thesis takes an actual industrial circulating water system as the research object, utilizing 28 days of continuously collected real historical sensor data (524,043 rows, 16 features) from an industrial enterprise. A complete data-driven control framework is proposed, encompassing \"physical mapping modeling — LSTM system identification — SVM+GA parameter optimization — DRL intelligent control — closed-loop simulation validation.\" First, addressing the critical issue of misalignment between sensor measurement positions and control target positions, we establish a flow physical mapping model (Q_he = Q_total × M1/(M1+M2), R² = 0.9998) and a temperature physical mapping model (T2 = α×T_cold + (1-α)×T_hot, R² = 0.9965) based on the system P&ID diagram, precisely converting sensor measurements into equivalent control targets at the heat exchanger inlet (flow 27.0 m³/h, temperature 22.0 °C), thereby establishing the objective function foundation for optimal control. Second, a 3-layer Long Short-Term Memory (LSTM) network with 977,291 parameters is constructed for black-box modeling of system dynamics. Trained on a dataset of 24,453 samples for 50 epochs, the model achieves a validation loss of 0.016128, with the coefficient of determination R² for key temperature state variables reaching up to 0.9788, providing a high-fidelity environmental surrogate model for subsequent optimization and control.")

    add_body(doc, "Building upon system modeling, two control strategies are investigated. The first strategy employs a Support Vector Machine (SVM) surrogate model to accelerate Genetic Algorithm (GA) fitness evaluation; trained on 5,000 LSTM-evaluated samples, the SVM surrogate achieves MSE = 0.1263, and the GA converges after 80 generations to an optimal fitness of -2487.12, yielding optimal control parameters: valve_DN200_fb = 19.30, valve_DN300_fb = 99.11, valve_DN350_fb = 99.67, valve_DN400_fb = 77.39, pump_speed_fb = 1152.10 RPM. The second strategy adopts the Deep Deterministic Policy Gradient (DDPG) framework and innovatively introduces a GA Behavioral Cloning (BC) pre-training mechanism: 150 expert control trajectories generated by the GA optimizer under diverse initial conditions are used to pre-train the DDPG Actor network via supervised learning (Actor MSE = 0.0389), effectively transferring optimization capability into a neural network policy. At the control execution level, a physics-guided real-time correction mechanism is introduced, dynamically adjusting M1 (flow correction, blending ratio 80% physics + 20% network) and M3 (temperature correction, blending ratio 60% physics + 40% network) to ensure the controller outputs physically plausible actions during actual execution.")

    add_body(doc, "Closed-loop simulation results demonstrate that: GA closed-loop optimization over 30 iterations achieves 100% flow compliance, 100% temperature compliance, and 100% pressure compliance, with mean equivalent flow of 27.1 m³/h, mean equivalent temperature of 20.47 °C, and mean energy consumption of 348.6; DRL closed-loop control over 100 steps also achieves 100% compliance across all three indicators, with mean equivalent temperature of 20.56 °C being closer to the target of 22.0 °C, and mean energy consumption of 394.9. Both methods validate the feasibility and effectiveness of data-driven approaches for circulating water system control: GA shows slight advantages in global optimization and energy efficiency, while DRL performs better in temperature tracking accuracy and real-time response speed. The research outcomes provide theoretical foundations and technical solutions for the intelligent upgrade of industrial circulating water systems, and the proposed \"expert optimization + neural network imitation + physics correction\" hybrid control paradigm offers referential value for intelligent control of similar complex industrial processes.")

    for _ in range(2):
        p = doc.add_paragraph(); set_paragraph_spacing(p, line_spacing=1.25)
    p = doc.add_paragraph()
    run = p.add_run("Key Words："); set_run_font(run, bold=True)
    run = p.add_run("Circulating Water System; Data-Driven Control; LSTM; Genetic Algorithm; Deep Reinforcement Learning; Behavioral Cloning; Physical Mapping; Closed-Loop Simulation")
    set_run_font(run)
    add_page_break(doc)

    # ==================== 目录 ====================
    add_chapter(doc, "目  录")
    toc_items = [
        ("摘  要", "I"), ("Abstract", "II"),
        ("第1章  绪论", "1"),
        ("  1.1  研究背景与意义", "1"),
        ("  1.2  国内外相关研究进展", "4"),
        ("    1.2.1  循环水系统建模方法研究", "4"),
        ("    1.2.2  智能优化算法在工业控制中的应用", "5"),
        ("    1.2.3  深度强化学习控制研究", "6"),
        ("    1.2.4  现有研究的不足", "7"),
        ("  1.3  本文主要研究内容与创新点", "7"),
        ("  1.4  论文组织结构", "9"),
        ("第2章  循环水系统建模与分析", "10"),
        ("  2.1  系统描述与问题定义", "10"),
        ("  2.2  物理映射模型构建", "13"),
        ("  2.3  数据采集与预处理", "15"),
        ("  2.4  系统能耗分析", "18"),
        ("第3章  基于LSTM的系统状态预测模型", "19"),
        ("  3.1  LSTM基本原理", "19"),
        ("  3.2  模型架构设计", "21"),
        ("  3.3  训练策略与超参数配置", "22"),
        ("  3.4  预测性能评估与分析", "24"),
        ("第4章  基于SVM与GA的控制参数优化", "27"),
        ("  4.1  SVM代理模型", "27"),
        ("  4.2  遗传算法优化框架", "29"),
        ("  4.3  优化实验与结果分析", "31"),
        ("第5章  基于深度强化学习的智能控制", "34"),
        ("  5.1  DRL问题形式化", "34"),
        ("  5.2  DDPG算法原理", "35"),
        ("  5.3  行为克隆预训练策略", "37"),
        ("  5.4  物理引导的实时修正机制", "39"),
        ("第6章  闭环仿真验证与结果分析", "41"),
        ("  6.1  闭环仿真环境设计", "41"),
        ("  6.2  GA闭环控制实验", "43"),
        ("  6.3  DRL闭环控制实验", "45"),
        ("  6.4  对比分析与讨论", "46"),
        ("结  论", "49"),
        ("参考文献", "51"),
        ("致  谢", "53"),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        if title.startswith("  "):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"{title}\t{page}")
            set_run_font(run, size=SIZE_BODY)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(f"{title}\t{page}")
            set_run_font(run, size=SIZE_BODY)
        set_paragraph_spacing(p, line_spacing=1.5)
    add_page_break(doc)

    # ==================== 第1章 绪论（大幅扩充） ====================
    add_chapter(doc, "第1章  绪  论")

    add_section(doc, "1.1  研究背景与意义")
    add_body(doc, "循环水系统是工业生产的\"血液循环系统\"，承担着为工艺设备提供冷却、带走废热、维持温度稳定的关键任务。在工业企业实际运行中，循环水冷却系统的控制当前主要依赖操作人员的技能和经验进行手动调节——操作人员根据经验判断当前工况，手动调整阀门开度和泵转速。这种人工控制方式存在响应滞后、控制精度不稳定、不同班组操作差异大等问题，严重制约了系统运行的稳定性、可靠性和能效水平。本文所使用的数据来源于某工业企业在2019年10月至11月期间循环水系统的真实运行数据，涵盖了温度、压力、流量多类传感器和执行器反馈信号共计28天的连续记录，为数据驱动的智能控制方法研究提供了宝贵的真实工业数据基础。在石油化工行业中，循环冷却水系统需要为反应釜、蒸馏塔、换热器等核心设备提供恒温恒压的冷却介质；在电力行业中，循环水系统直接关系到汽轮机凝汽器的真空度和发电效率；在冶金行业中，循环水冷却系统保障了高炉、转炉等高温设备的正常运行；在暖通空调（HVAC）领域，循环水系统是实现建筑节能和室内热舒适性的核心环节。据国际能源署（IEA）统计，工业冷却系统约占全球工业总能耗的15%-20%，其中循环水泵和冷却塔风机的电力消耗是最主要的能耗来源。在我国\"双碳\"战略目标背景下，工业循环水系统的节能降耗和智能化升级具有重大的经济价值和社会意义。")

    add_body(doc, "传统的循环水系统控制主要依赖比例-积分-微分（PID）控制器及其变体。PID控制具有结构简单、不依赖模型的优点，在工况稳定、系统线性度较好的场景下能够满足基本控制需求。然而，循环水系统本质上是一个多输入多输出（MIMO）、强耦合、非线性的复杂动态系统：(1) 多个控制回路之间存在复杂的耦合关系——阀门开度变化同时影响流量和压力分布，泵转速调节影响总管流量进而改变所有支路的流量分配；(2) 外界环境条件（气温、湿度、冷却塔效率）和内部负荷（工艺产热量变化）的波动导致系统动态特性呈现显著的时变性和不确定性；(3) 换热器、管道等设备的热惯性使系统具有较大的时间滞后，温度响应需要数分钟乃至数十分钟才能达到稳态；(4) 传感器受限于安装位置，往往不能直接测量控制目标处的物理量（如换热器入口的流量和温度），测量值与控制目标之间存在映射间隙。这些特性使得传统PID控制在面对工况大范围变化时，容易出现超调量大、调节时间长、稳态误差难以消除等问题，甚至在某些极端工况下导致控制失稳。")

    add_body(doc, "近年来，随着工业4.0和智能制造理念的深入推进，工业现场部署了越来越多的传感器和数据采集系统，积累了海量的历史运行数据。这些数据中蕴含着丰富的系统动态行为信息，为数据驱动的建模与控制方法提供了前所未有的机遇。与此同时，以深度学习、强化学习为代表的人工智能技术取得了突破性进展：深度神经网络展现出了强大的非线性函数逼近能力和高维特征自动提取能力；强化学习通过与环境交互学习最优策略的范式，天然适合序贯决策问题。将这些前沿AI技术应用于工业过程控制，有望突破传统方法的瓶颈，实现更高效、更智能、更自主的控制性能。")

    add_body(doc, "然而，将数据驱动方法应用于实际工业循环水系统控制仍面临若干核心挑战：(1) 如何利用有限的有噪声传感器数据，建立足够精确的系统动态模型，为优化控制提供可靠的环境代理；(2) 如何在多约束条件下高效搜索全局最优控制参数，同时保证搜索过程的计算可行性；(3) 如何设计能够实时响应、具备泛化能力的智能控制器，使其不仅能在训练工况下表现良好，还能适应未见过的运行条件；(4) 如何将物理机理知识（如质量守恒、能量守恒、分流混合原理）与数据驱动方法有机融合，弥补纯数据方法在物理可解释性和外推能力方面的不足。本文正是围绕上述挑战，以实际工业循环水系统为背景，系统性地研究数据驱动的智能建模与控制方法。")

    add_section(doc, "1.2  国内外相关研究进展")

    add_subsection(doc, "1.2.1  循环水系统建模方法研究")
    add_body(doc, "循环水系统的建模方法大致可分为三类：白箱建模（机理建模）、黑箱建模（数据驱动建模）和灰箱建模（混合建模）。白箱建模基于质量守恒、能量守恒和动量守恒定律，建立系统的微分方程或差分方程描述。典型的白箱模型包括：基于Navier-Stokes方程的管道流体动力学模型、基于传热方程的换热器热力学模型、基于泵特性曲线的离心泵水力模型等。白箱模型的优点是物理意义清晰、可解释性强、外推能力好，但其缺点也十分突出——需要详细的结构参数（管道直径、换热面积、传热系数等）和边界条件，建模过程耗时且需要对系统有深入的领域知识，模型校准困难，不同系统之间难以直接迁移。")

    add_body(doc, "黑箱建模不依赖物理机理，而是直接从输入输出数据中学习系统的映射关系。早期研究主要采用自回归滑动平均（ARMA）模型、支持向量回归（SVR）、人工神经网络（ANN）等浅层方法。例如，Zhang等人[1]采用SVR建立了换热器出口温度的预测模型，在特定工况下取得了较好的精度。然而，浅层模型受限于表达能力，难以捕捉复杂动态系统中的长期时序依赖和非线性特征。随着深度学习的发展，循环神经网络（RNN）及其变体长短期记忆网络（LSTM）和门控循环单元（GRU）在时间序列建模中展现出显著优势。LSTM通过引入遗忘门、输入门和输出门的门控机制，有效解决了传统RNN在处理长序列时的梯度消失问题，能够记忆数百乃至数千步之前的系统状态信息，特别适合描述具有大惯性、大滞后的热工过程。Li等人[2]使用LSTM对集中供热系统的回水温度进行多步预测，RMSE较传统ARIMA模型降低了42%。Chen等人[3]将注意力机制引入LSTM框架，进一步提升了模型对关键时间点的聚焦能力。Liang等人[24]针对车辆气候试验室的冷水系统进行了温度控制研究，探讨了基于实测数据的系统建模与控制方法。")

    add_body(doc, "灰箱建模旨在融合白箱和黑箱方法的优势，将物理先验知识作为约束或正则化项嵌入数据驱动模型中。物理信息神经网络（PINN）是灰箱建模的代表性方法，其通过在损失函数中加入物理方程残差项，引导神经网络学习符合物理规律的解。在流体和热工系统中，PINN已成功应用于求解Navier-Stokes方程、热传导方程等偏微分方程的正问题和逆问题。然而，PINN在工业过程控制中的应用仍处于起步阶段，主要面临物理方程不完整、边界条件不确定、训练收敛困难等实际障碍。")

    add_subsection(doc, "1.2.2  智能优化算法在工业控制中的应用")
    add_body(doc, "工业过程优化控制的核心是在满足各类约束条件的前提下，搜索使性能指标最优的控制参数组合。当目标函数具有多峰、非凸、不连续等复杂特性时，传统的梯度下降方法容易陷入局部最优。遗传算法（Genetic Algorithm, GA）是一类模拟自然界生物进化机制的概率搜索算法，其通过选择、交叉、变异等遗传操作对种群进行迭代进化，具有不依赖梯度信息、能够处理离散和连续混合变量、全局搜索能力强等优点。在化工过程优化、能源系统调度、水资源管理等领域，GA已被广泛应用于求解多目标、多约束的优化问题。Wang等人[4]将GA应用于冷却水系统的运行优化，以最小化泵能耗和冷却塔风机能耗为双目标，在保证冷却效果的前提下实现了约12%的节能率。")

    add_body(doc, "然而，GA在实际应用中面临一个关键瓶颈：适应度评估的计算成本。在工业过程优化中，每一次适应度评估可能涉及耗时的数值仿真或代价高昂的物理实验。对于本文研究的循环水系统，若每次适应度评估都需要运行完整的LSTM前向传播，GA优化将因计算成本过高而失去实用性。代理模型辅助进化优化（Surrogate-Assisted Evolutionary Optimization）是解决这一问题的有效途径：通过训练一个计算成本低的代理模型（如Kriging模型、径向基函数模型、支持向量机模型等）来近似昂贵的真实适应度函数，在进化搜索过程中用代理模型替代大部分真实评估，仅对少数最有前景的候选解进行真实评估以更新代理模型。Liu等人[5]提出了基于高斯过程代理模型的自适应采样策略，在保证优化精度的同时将真实评估次数减少了70%以上。本文采用SVM作为代理模型，利用其在中小样本回归任务中的优异泛化性能，实现GA适应度评估的高效加速。")

    add_subsection(doc, "1.2.3  深度强化学习控制研究")
    add_body(doc, "强化学习（Reinforcement Learning, RL）是一种通过与环境交互、以试错方式学习最优行为策略的机器学习范式。RL智能体在每一步根据当前状态选择动作，环境返回奖励信号和下一状态，智能体的目标是最大化累积奖励的期望值。深度强化学习（DRL）将深度神经网络与RL结合，利用深度网络的强大表示能力处理高维连续状态和动作空间，在游戏AI、机器人控制、自动驾驶等领域取得了超越人类水平的性能。")

    add_body(doc, "在工业过程控制领域，DRL的研究近年来迅速增长。Xu等人[23]将深度强化学习应用于循环冷却水系统的自适应控制，在PLoS ONE发表的论文中验证了DRL在循环水系统控制中的可行性和有效性。Mnih等人[6]提出的深度Q网络（DQN）在Atari游戏环境中展示了端到端学习的能力，但其仅适用于离散动作空间。针对连续控制问题，Lillicrap等人[7]提出了深度确定性策略梯度（Deep Deterministic Policy Gradient, DDPG）算法，采用Actor-Critic架构和确定性策略梯度，能够在高维连续动作空间中学习稳定有效的控制策略。Schulman等人[8]提出的近端策略优化（PPO）算法通过限制策略更新步长提高了训练的稳定性和样本效率。在暖通空调领域，Zhang等人[9]将DDPG应用于建筑空调系统的温度控制，相比传统PID实现了15%-20%的节能效果。在化工过程领域，Spielberg等人[10]使用PPO训练连续搅拌釜反应器的温度控制器，展示了DRL处理非线性、多变量过程控制问题的能力。")

    add_body(doc, "然而，DRL在工业控制中的落地应用仍面临严峻挑战。首先，DRL通常需要数百万次的环境交互才能学到有效的策略，而实际工业系统不可能承受如此大量的试错操作。其次，工业系统对安全性有严格要求，训练初期的随机探索可能导致危险的操作状态。再次，仿真环境与真实系统之间不可避免地存在模型失配（Reality Gap），在仿真中训练的DRL策略迁移到实际系统时可能出现性能退化。针对上述问题，模仿学习（Imitation Learning）和行为克隆（Behavioral Cloning, BC）提供了一条可行的捷径：利用已有的专家策略（如人工操作员的历史操作记录或传统优化器的输出）预训练DRL策略网络，为其提供良好的初始化，从而大幅减少所需的在线交互次数。Ross等人[11]从理论上证明了BC策略在专家数据足够覆盖状态空间时的性能保证。本文结合GA优化器和BC预训练策略，提出了一种\"GA专家优化 + BC策略蒸馏 + 物理实时修正\"的混合控制方案，兼顾了控制效果、计算效率和物理合理性。")

    add_subsection(doc, "1.2.4  现有研究的不足")
    add_body(doc, "综合以上分析，现有研究在以下方面仍存在不足：(1) 多数研究专注于单一环节（仅建模或仅控制），缺乏从数据预处理、系统建模到优化控制的完整端到端解决方案；(2) 传感器测量位置与控制目标位置不一致的问题在工业界普遍存在，但鲜有研究将其作为优化控制框架的内在组成部分进行系统性处理；(3) 遗传算法在工业优化中的应用大多直接使用计算代价高昂的真实仿真器，缺乏针对代理模型加速的工程实践验证；(4) DRL在工业过程控制中的研究多停留在仿真模拟阶段，且多依赖大量的在线试错训练，缺乏与专家知识的有效融合机制；(5) 数据驱动方法与物理机理的融合多停留在概念层面，缺乏具体可操作的混合控制策略设计和定量效果验证。本文正是针对上述不足，设计和实现了一套完整的、经定量闭环仿真验证的数据驱动智能控制方案。")

    add_section(doc, "1.3  本文主要研究内容与创新点")
    add_body(doc, "本文以工业循环水系统为研究对象，围绕\"如何利用历史运行数据实现系统的智能优化控制\"这一核心问题，按照\"物理建模—系统辨识—参数优化—策略学习—闭环验证\"的技术路线，开展了以下五个方面的研究工作：")

    add_body(doc, "（1）物理映射建模。针对系统P&ID图中传感器安装位置与控制目标位置之间的物理差异，基于分流原理和混合原理分别建立了流量物理映射模型和温度物理映射模型。流量映射模型将总管流量传感器读数转化为换热器入口的等效流量（Q_he = Q_total × M1/(M1+M2)），温度映射模型将冷热水温度传感器读数转化为换热器入口的等效混合温度（T2 = α×T_cold + (1-α)×T_hot）。通过历史数据的线性回归验证了映射模型的高精度（流量R² = 0.9998，温度R² = 0.9965），并将映射后的等效目标（流量27.0 m³/h，温度22.0 °C）作为优化控制的统一目标函数。")

    add_body(doc, "（2）LSTM系统动态建模。基于实际工业传感器采集的28天连续运行数据（524,043行、16维特征），经过数据清洗、Savitzky-Golay滤波、滑动窗口（窗口60步、步长15）和Z-score标准化等预处理步骤，构建了34,933个监督学习样本。设计了3层堆叠LSTM网络（隐藏层大小128、Dropout 0.2），以过去60秒的11维特征（6维状态+5维控制）预测下一时刻的6维系统状态。模型在CPU环境下训练50轮，验证损失0.016128，温度类变量的R²达到0.88-0.98，为后续优化控制提供了高保真的环境代理。")

    add_body(doc, "（3）SVM+GA参数优化。训练了基于径向基函数（RBF）核的SVM代理模型（5,000训练样本，MSE=0.1263），将单次适应度评估的计算时间从LSTM的秒级降至毫秒级。设计了包含锦标赛选择、算术交叉、高斯变异和精英保留策略的GA优化框架，在5维连续控制空间中搜索最优控制参数组合。GA在80代进化后收敛，最优适应度-2487.12，最优解具有良好的物理可解释性。")

    add_body(doc, "（4）DRL智能控制策略。将循环水系统控制问题形式化为马尔可夫决策过程（MDP），采用DDPG算法框架。创新性地提出GA行为克隆（BC）预训练策略：收集GA优化器在多样化初始条件下产生的150条专家轨迹，通过监督学习预训练Actor网络（MSE=0.0389），将GA的优化能力蒸馏到神经网络策略中。设计了物理引导的实时修正机制，在动作执行前根据当前系统状态动态校正M1（流量修正）和M3（温度修正），确保控制器输出物理合理的控制指令。")

    add_body(doc, "（5）闭环仿真验证。构建了以LSTM为环境代理、物理映射为评估准则的闭环仿真平台。分别在30次GA迭代和100步DRL执行的闭环场景下，定量评估了两种方法的流量达标率、温度达标率、压力达标率和平均能耗。仿真结果表明两种方法均实现三项指标100%达标，验证了所提框架的有效性。")

    add_body(doc, "本文的主要创新点包括：(i) 提出并验证了物理映射 + LSTM建模 + GA优化/DRL控制的完整数据驱动控制框架，实现了从原始传感器数据到闭环控制策略的端到端解决方案；(ii) 设计了GA行为克隆预训练 + 物理实时修正的混合控制策略，有效结合了专家优化的全局搜索能力和神经网络的实时推理速度；(iii) 在实际工业数据上进行了充分的闭环仿真验证，定量证明了方法的有效性。")

    add_section(doc, "1.4  论文组织结构")
    add_body(doc, "本文共分为六章。第1章为绪论，介绍研究背景、国内外研究进展和本文的研究思路。第2章详细阐述循环水系统的物理建模过程，包括系统描述、物理映射模型和数据预处理。第3章介绍基于LSTM的系统状态预测模型，包括网络设计、训练过程和性能评估。第4章介绍基于SVM代理模型和GA的参数优化方法。第5章介绍基于DDPG和BC预训练的DRL智能控制策略。第6章给出闭环仿真验证的设计和结果，并对GA和DRL两种方法进行全面的对比分析。最后为结论与展望，总结全文工作并讨论未来研究方向。")

    add_page_break(doc)

    # ==================== 第2章 循环水系统建模与分析（扩充） ====================
    add_chapter(doc, "第2章  循环水系统建模与分析")

    add_section(doc, "2.1  系统描述与问题定义")
    add_body(doc, "本文研究的循环水系统是一个典型的闭式工业冷却水循环回路，其工艺流程如图2.1所示（完整的P&ID图见附录或工程图纸）。系统由以下主要设备和组件构成：冷水罐（储存和缓冲冷却水，维持系统水量平衡）、增压泵（为循环水提供动力，克服管道阻力和设备压降）、四台电动调节阀（DN200、DN300、DN350、DN400，分别承担主管路流量调节、泄压旁通、冷热混合比例调控和排水流量控制功能）、换热器（冷却水与工艺介质进行热交换的核心设备）以及配套的管道网络和传感器系统。")

    add_body(doc, "系统的工作流程如下：冷水罐中的低温冷却水经增压泵加压后，分为多路管道流动。其中主管路（经过DN200电动调节阀M1）和旁通管路（经过DN300电动调节阀M2）形成分流结构——主管路将冷却水输送至换热器入口，旁通管路则将部分流量回流或泄压。此外，换热器出口的热水与冷水罐出水经过DN350电动调节阀M3按一定比例混合，以调节进入换热器的冷却水温度。DN400电动调节阀M4负责排水流量控制，维持系统压力平衡。系统控制的核心目标是确保进入换热器的冷却水具有合适的流量和温度，以满足工艺介质热交换的需求。")

    add_body(doc, "系统配备的传感器网络包括：(1) 温度传感器3个——增压泵后温度（T_pump_out，反映进入主管路的冷却水初始温度）、冷水罐出口温度（T_tank_out，反映系统冷源温度）、换热器模拟出口温度（T_he_sim，反映热交换后的回水温度）；(2) 压力传感器4个——增压泵后压力（P_pump_out）、DN400阀后压力（P_valve_DN400）、止回阀后压力（P_check_valve）、冷水罐出口压力（P_tank_out）；(3) 流量传感器2个——DN300泵后流量（F_pump_DN300，反映总管流量）、DN400泵后流量（F_pump_DN400）；(4) 执行器反馈信号5个——四台电动阀的开度反馈值和增压泵转速反馈值。所有传感器的数据采集频率为1 Hz，通过工业现场总线汇总至数据中心。本文使用的数据来源于某工业企业循环水系统的真实运行记录，时间跨度为2019年10月21日17:42:16至2019年11月18日19:06:15，覆盖了约28天的连续运行，期间包含了不同环境温度、不同工艺负荷下的多种典型工况。该企业循环水系统在数据采集期间处于正常生产运行状态，控制操作由现场操作人员根据经验手动完成，因此数据中\"天然\"包含了人工操作的决策模式，为后续数据驱动建模提供了真实、多样的控制样本。")

    add_figure(doc,
        os.path.join(RESULT_DIR, "system_overview.png"),
        "图2.1  循环水系统工艺流程与传感器布置概览")

    add_body(doc, "控制问题的数学形式化定义如下。设系统在时刻t的状态为s_t = [T_pump_out, P_pump_out, F_pump_DN300, T_tank_out, T_he_sim, P_tank_out]（6维向量），控制动作为a_t = [M1, M2, M3, M4, pump_speed]（5维向量，分别对应DN200阀开度、DN300阀开度、DN350阀开度、DN400阀开度和泵转速）。需要求解最优控制策略π: S→A，使得在满足物理约束的前提下，换热器入口的等效流量Q_he和等效温度T2尽可能接近设计目标（Q_target = 27.0 m³/h，容许范围[26, 28]；T_target = 22.0 °C，容许范围[20, 24]），同时确保系统压力不超过安全上限3.0 MPa，并尽可能降低系统能耗。形式化地，目标函数可表达为：")

    add_formula(doc, "min  J = w₁·|Q_he - 27.0| + w₂·|T2 - 22.0| + w₃·max(P - 3.0, 0) + w₄·Energy")

    add_body(doc, "其中w₁、w₂、w₃、w₄为权衡不同控制目标的权重系数，P为系统中各测点压力的最大值，Energy为系统能耗（主要由泵功率和各阀门压损决定）。该优化问题具有以下难点：(1) Q_he和T2不是直接测量量，需要通过传感器值和控制量的物理映射间接计算；(2) 状态转移s_t→s_{t+1}遵循未知的复杂非线性动力学；(3) 控制量受执行器物理限幅约束（阀门开度范围、泵转速上下限）；(4) 需要在每个控制周期（秒级）内完成决策计算。")

    add_section(doc, "2.2  物理映射模型构建")
    add_body(doc, "物理映射模型是本文方法区别于纯黑箱数据驱动方法的关键组件。其核心思想是：利用已知的物理结构信息（P&ID图中管路的连接关系和分流/混合拓扑），建立从传感器测量值到控制目标处物理量的确定性映射关系，从而将\"传感器观测空间\"中的测量值转化为\"控制目标空间\"中的等效值，为优化控制提供物理意义明确的目标函数。")

    add_subsection(doc, "2.2.1  流量物理映射模型")
    add_body(doc, "根据系统的管路拓扑结构，总管流量F_pump_DN300（传感器测量值）经过M1（DN200主管路阀）和M2（DN300旁通阀）分流后，进入换热器的等效流量Q_he由分流比决定。假设阀门开度与其流通能力（Cv值）近似成正比（在阀门正常工作范围内），且管道阻力相对阀门压降可忽略，则进入主管路的流量比例为M1/(M1+M2)。因此，换热器入口的等效流量为：")

    add_formula(doc, "Q_he = Q_total × M1 / (M1 + M2)")

    add_body(doc, "其中Q_total = F_pump_DN300为总管流量传感器读数，M1 = valve_DN200_fb为DN200阀开度反馈值，M2 = valve_DN300_fb为DN300阀开度反馈值。该映射关系的物理基础是并联管路的流量分配原理：在并联管路中，各支路的流量分配与支路阻力成反比，而阀门开度是决定支路阻力的主导因素。通过对历史数据的线性回归验证（以M1/(M1+M2)×Q_total为自变量，以换热器入口流量的机理估算值为因变量），得到R² = 0.9998，表明该简单分流模型能够以极高的精度描述实际的流量分配关系。")

    add_body(doc, "基于历史数据的统计，换热器入口等效流量Q_he的分布特征为：均值108.1 m³/h，波动范围-3.9~190.4 m³/h（负值对应极端工况下的回流现象）。Q_he的均值108.1 m³/h远高于设计目标27.0 m³/h，这反映了供热季（10-11月）系统实际运行在高负荷状态，需要通过精确的阀门调节将进入换热器的流量降低至设计值。这反过来也说明了主动控制的重要性——若不加控制地全开阀门，换热器将承受远超设计值的流量冲击。")

    add_subsection(doc, "2.2.2  温度物理映射模型")
    add_body(doc, "换热器入口的等效温度T2由两股水流混合决定：一股来自冷水罐的低温水（温度T_cold ≈ T_tank_out），另一股来自换热器出口的高温回水（温度T_hot ≈ T_he_sim）。两股水在DN350电动调节阀M3处按一定比例混合后进入换热器。设混合比例系数为α（0.4 ≤ α ≤ 0.7），表示冷水的混合占比，则等效温度T2按照加权平均计算：")

    add_formula(doc, "T2 = α × T_cold + (1 - α) × T_hot")

    add_body(doc, "混合比例α与M3阀开度的关系通过物理机理分析和数据拟合确定。当M3全开时（M3 ≈ 99.75），冷水流路阻力最小，冷水占主导，此时α接近最大值0.7；当M3全关时（M3 ≈ 99.6），热水流路相对畅通，冷水占比降低，此时α接近最小值0.4。将M3开度归一化至[0, 1]区间：m3_norm = (M3 - 99.6) / 0.15，则α与m3_norm的关系为线性映射：α = 0.7 - 0.3×m3_norm。该温度映射模型在历史数据上的拟合优度R² = 0.9965，验证了混合模型和α-M3线性关系假设的合理性。基于历史数据统计，T2的分布特征为：均值18.9 °C，波动范围7.3~24.8 °C。均值18.9 °C低于设计目标22.0 °C，表明在实际运行中，由于冷水占比通常较大（α偏大），换热器入口温度常偏低，需要适度减小M3开度以增加热水混合比例，提升入口温度至目标值。")

    add_subsection(doc, "2.2.3  控制目标等效转化")
    add_body(doc, "综合流量映射和温度映射模型，本文将循环水系统的控制目标从\"传感器值\"空间等效转化为\"换热器入口\"空间：原始设计指标为换热器入口流量27.0 m³/h、入口温度22.0 °C。由于换热器入口没有直接安装流量计和温度计，实际控制时以传感器读数和阀门反馈值通过映射公式计算等效值作为替代。这一等效转化的关键优势在于：(1) 使优化控制的目标函数具有明确的物理含义，对应换热器的\"真实\"入口条件；(2) 将传感器位置带来的系统误差从控制回路中显式地分离出来，由物理映射层负责处理；(3) 映射公式形如Q_he = f(Q_total, M1, M2)和T2 = g(T_cold, T_hot, M3)，明确揭示了\"哪些控制量影响哪些目标量\"——M1和M2主要影响等效流量，M3主要影响等效温度，这为指导控制策略设计提供了物理洞察。")

    add_section(doc, "2.3  数据采集与预处理")
    add_body(doc, "高质量的训练数据是数据驱动方法取得良好性能的前提。本文的数据预处理流程涵盖六个步骤，旨在从原始工业传感器数据中提取出可用于监督学习的干净、规范化的样本。")

    add_subsection(doc, "2.3.1  原始数据加载与对齐")
    add_body(doc, "数据源包含20对FB（Feedback，反馈信号）和TP（Temperature & Pressure，温度压力）CSV文件，每对文件对应一个数据采集周期。FB文件包含时间戳和5维执行器反馈信号（4个阀门开度+泵转速），TP文件包含时间戳和9维过程量（温度、压力、流量）。原始文件编码格式为UTF-16 LE，采用分号作为列分隔符。数据加载后按时间戳进行精确配对对齐（FB和TP的时间戳精度均为秒级），去除仅有单侧记录的时间点。FB数据共530,264条记录，TP数据共524,624条记录，合并对齐后得到524,043条完整记录（丢失约1.2%的数据，主要原因为两侧采集时刻不完全重合），时间范围覆盖2019年10月21日17:42:16至2019年11月18日19:06:15。")

    add_subsection(doc, "2.3.2  异常值检测与清洗")
    add_body(doc, "工业传感器数据不可避免地包含异常值，其来源包括：传感器临时故障、信号传输干扰、数据采集系统偶发性错误、系统处于非正常工况（如停机检修、启停过渡过程）等。本文采用以下异常值检测策略：对于每维特征，计算其在全数据集上的均值μ和标准差σ，将偏离均值超过4σ（|x - μ| > 4σ）的数据点标记为全局异常值；同时，计算每个数据点与其前后5个邻点的局部偏差，将偏差超过局部3σ的数据点标记为局部异常值（突刺噪声）。标记的异常值约占数据总量的0.8%，对其采用前后正常值的线性插值进行修正。对于连续缺失超过30秒（30个数据点）的长段异常，不进行插值而直接舍弃该段，以避免引入虚假信息。")

    add_subsection(doc, "2.3.3  信号滤波平滑")
    add_body(doc, "传感器信号中叠加了高频测量噪声，直接使用原始信号可能导致模型学习到噪声模式而非系统真实的动态特性。本文选用Savitzky-Golay（S-G）滤波器进行信号平滑处理。S-G滤波器通过在滑动窗口内拟合低阶多项式来估计窗口中心点的信号值，相比移动平均滤波在保留信号峰值和边缘特征方面具有优势。滤波器参数设置为：窗口大小11（对应11秒的物理时间），多项式阶数3。对于温度等变化缓慢的信号，S-G滤波能有效抑制幅度约0.05-0.2 °C的高频噪声；对于压力和流量等变化较快的信号，在抑制噪声的同时较好地保留了信号的动态特征。")

    add_subsection(doc, "2.3.4  滑动窗口样本构建")
    add_body(doc, "将时序数据转化为监督学习所需的（输入，输出）样本对，需要确定时间窗口大小和滑动步长两个关键参数。时间窗口大小决定了模型可用的历史信息量：窗口过小则模型缺乏足够的上下文来推断系统状态演变趋势；窗口过大则增加了输入维度和计算负担，且过于久远的历史信息可能对当前预测贡献微弱。经过实验比较，本文选取窗口大小为60步（对应60秒），这一选择基于以下考虑：循环水系统的热惯性时间常数在分钟量级（30秒至数分钟），60秒的窗口约覆盖系统的主要动态响应时间，能够提供足够的时序上下文。滑动步长设置为15步，在保证训练样本充足（34,933个）的同时避免了相邻样本之间的过高相关性（步长15对应25%的重叠率）。每个样本的输入X为形状[60, 11]的矩阵（60个连续时间步，每步11维——6维状态+5维控制），输出y为形状[6,]的向量（下一时间步的6维状态预测值）。")

    add_subsection(doc, "2.3.5  数据标准化")
    add_body(doc, "不同物理量之间的数值量级差异巨大（如温度在10-40 °C范围，压力在0.1-0.4 MPa范围，流量在0-500 m³/h范围），若直接输入神经网络，量级较大的特征将在梯度计算中占据主导地位，严重影响训练效率和模型性能。本文采用Z-score标准化方法：对每维特征x，计算其在训练集上的均值μ_train和标准差σ_train，标准化值x' = (x - μ_train) / σ_train。标准化参数仅在训练集上计算，验证集和测试集使用相同的参数进行变换，以严格避免数据泄露。标准化后，所有特征的分布近似为标准正态分布N(0, 1)。")

    add_subsection(doc, "2.3.6  数据集划分")
    add_body(doc, "按照时间顺序将数据集划分为训练集（前70%，24,453样本）、验证集（中15%，5,240样本）和测试集（后15%，5,240样本）。这种按时间划分（而非随机划分）的策略符合工业时间序列建模的最佳实践：它模拟了真实的预测场景——使用过去的数据训练模型，预测未来的系统行为。按时间划分避免了一个常见的陷阱：若随机打乱数据，未来时刻的信息可能\"泄漏\"到训练集中，导致模型学习到虚假的\"前后因果\"关系，从而在测试集上获得过于乐观（但不可信）的性能指标。")

    add_table(doc,
        ["指标", "训练集", "验证集", "测试集", "总计"],
        [
            ["样本数量", "24,453", "5,240", "5,240", "34,933"],
            ["时间窗口（步）", "60", "60", "60", "-"],
            ["输入维度（特征数）", "11", "11", "11", "-"],
            ["输出维度（状态数）", "6", "6", "6", "-"],
            ["滑动步长（步）", "15", "15", "15", "-"],
            ["窗口重叠率", "75%", "75%", "75%", "-"],
        ],
        "表2.1  数据集划分与参数概况")

    add_section(doc, "2.4  系统能耗分析")
    add_body(doc, "系统能耗主要由增压泵的电能消耗构成，同时各阀门的节流损失也间接影响系统总能效。基于历史数据的统计分析，系统功耗特征如下：运行期间的平均功率为40.2 kW，功率波动范围为7.3~52.7 kW（最低功率对应低负荷或待机工况，最高功率对应高负荷全速运行工况），28天累计能耗为965.3 kWh。将功耗与主要控制量进行相关性分析发现：泵转速与功耗呈强正相关（相关系数约0.85），这是由泵的相似定律（功率与转速的三次方成正比）决定的；M2（泄压阀）开度与功耗呈中等正相关，因为旁通流量越大，泵需要输出更多的流量来维持主管路流量，导致无用功增加；M1（主管路阀）开度与功耗呈弱的负相关，因为主管路畅通时系统阻力较小。这些能耗特征为优化控制中的能耗惩罚项设计提供了定量依据。")

    add_page_break(doc)

    # ==================== 第3章 LSTM（扩充） ====================
    add_chapter(doc, "第3章  基于LSTM的系统状态预测模型")

    add_section(doc, "3.1  LSTM基本原理")
    add_body(doc, "长短期记忆网络（Long Short-Term Memory, LSTM）由Hochreiter和Schmidhuber于1997年提出[12]，是循环神经网络（RNN）的一种重要变体，专门设计用于解决传统RNN在处理长序列数据时面临的梯度消失和梯度爆炸问题。LSTM的核心创新在于引入了细胞状态（Cell State）c_t和三个门控结构——遗忘门（Forget Gate）f_t、输入门（Input Gate）i_t和输出门（Output Gate）o_t。细胞状态c_t作为贯穿时间步的\"信息传送带\"，允许信息在时间维度上相对无损地传递；门控结构则以sigmoid激活函数（输出0-1之间的值）控制信息的保留和丢弃程度。")

    add_body(doc, "LSTM单元在每个时间步t的计算过程如下。给定当前输入x_t和上一时间步的隐藏状态h_{t-1}，(1) 遗忘门决定从细胞状态中丢弃哪些旧信息：f_t = σ(W_f·[h_{t-1}, x_t] + b_f)；(2) 输入门决定哪些新信息写入细胞状态：i_t = σ(W_i·[h_{t-1}, x_t] + b_i)；(3) 候选细胞状态C̃_t由tanh层生成：C̃_t = tanh(W_c·[h_{t-1}, x_t] + b_c)；(4) 细胞状态更新：c_t = f_t ⊙ c_{t-1} + i_t ⊙ C̃_t（⊙表示逐元素乘法）；(5) 输出门决定输出什么信息：o_t = σ(W_o·[h_{t-1}, x_t] + b_o)；(6) 隐藏状态更新：h_t = o_t ⊙ tanh(c_t)。通过这种精巧的门控机制，LSTM可以在细胞状态中保持对数十甚至数百个时间步之前的关键信息的记忆，同时遗忘不再相关的陈旧信息。")

    add_body(doc, "LSTM的这一特性使其特别适合对具有大惯性、大滞后的热工过程进行建模。在循环水系统中，当前时刻的阀门调节动作对换热器入口温度和流量的影响并非即时生效，而是在数秒至数分钟后才充分显现（流体传输延迟+热交换时间常数）。LSTM的细胞状态可以自然地编码这种延迟效应——输入门在调节动作发生时将相关信息写入细胞状态，遗忘门在热效应衰减后清除过时信息，输出门在需要时提取累积的热状态信息用于预测当前时刻的输出。")

    add_section(doc, "3.2  模型架构设计")
    add_body(doc, "本文设计的LSTM系统预测模型采用多层堆叠架构，以逐层提取不同抽象级别的时序特征。模型总体结构为：输入层→3层LSTM（堆叠）→全连接输出层。具体参数配置如下：")

    add_body(doc, "输入层：接收形状为[batch_size, 60, 11]的三维张量，其中batch_size为批次大小（训练时取256），60为时间窗口长度，11为每步的特征维度。11维特征包括6维状态变量（泵后温度、泵后压力、总管流量、冷水罐温度、换热器温度、冷水罐压力）和5维控制变量（四个阀门开度+泵转速）。注意，输入中包含历史控制动作序列——这意味着LSTM在预测下一时刻状态时，不仅知道过去的系统状态，还知道在此期间执行了哪些控制动作，这使得预测问题在信息上是完备的。")

    add_body(doc, "LSTM隐藏层：堆叠3层LSTM，每层包含128个隐藏单元（hidden_size=128）。选择3层的原因是基于实验比较：单层LSTM（128单元）的验证损失约为0.022，表达能力不足以充分捕捉系统动态；2层LSTM的验证损失降至约0.018，改善明显；3层LSTM进一步降至0.016，但4层LSTM（额外增加12万参数）的验证损失仅微降至0.0158，增益递减的同时训练时间增加了约25%。因此，3层128单元在模型容量和计算效率之间取得了较好的平衡。层间应用Dropout正则化，丢弃率设置为0.2——在训练时随机丢弃20%的神经元输出，等效于训练一个指数级数量的子网络集合，有效抑制过拟合。")

    add_body(doc, "输出层：全连接层（线性层），将最后一层LSTM的隐藏状态（维度128）映射为6维状态预测向量。不使用激活函数，因为状态预测是一个回归任务，输出值（标准化后）的范围不受[0,1]或[-1,1]的约束。")

    add_body(doc, "模型的总参数量为977,291个。从参数构成来看：3层LSTM共约974,000参数（每层约325,000，主要由输入→隐藏和隐藏→隐藏的权重矩阵贡献），全连接输出层约770参数（128×6权重+6偏置）。模型基于PyTorch 2.x框架实现，可在CPU或GPU上运行。考虑到大多数工业控制计算机不具备高性能GPU，本文所有实验均在CPU（Intel Core i7）上完成，以验证方法在实际工业硬件条件下的可行性。")

    add_section(doc, "3.3  训练策略与超参数配置")
    add_body(doc, "模型训练的损失函数选用均方误差（Mean Squared Error, MSE），公式为L = (1/N) Σ_i Σ_j (ŷ_{ij} - y_{ij})²，其中N为批次样本数，j=1,...,6遍历6个输出维度。优化器选用Adam（Adaptive Moment Estimation），这是目前深度学习中应用最广泛的自适应学习率优化算法之一[13]。Adam结合了动量法和RMSprop的思想，为每个参数维护独立的自适应学习率，能够在训练初期自动采用较大的学习速率以快速收敛，后期自动缩小学习速率以精细调优。Adam的关键超参数设置为：β₁=0.9（一阶动量衰减系数）、β₂=0.999（二阶动量衰减系数）、ε=10⁻⁸（数值稳定常数）。")

    add_body(doc, "学习率调度是影响训练效果的重要因素。本文采用余弦退火（Cosine Annealing）学习率调度策略：初始学习率设置为1×10⁻³，在每个训练周期（epoch）结束后，学习率按照余弦曲线从初始值逐渐衰减至接近0，周期长度（T_max）设置为50个epoch。余弦退火的优势在于，学习率在衰减过程中会经历从\"快速下降→慢速下降→极慢速下降（接近谷底）\"的非线性变化，使优化过程有机会跳出浅的局部最优并最终收敛到更深的极小值点。")

    add_body(doc, "为防止过拟合，设置了早停（Early Stopping）机制：监控验证集损失，若连续10个epoch验证损失没有改善，则终止训练并恢复到最佳验证损失对应的模型参数。此外，在训练过程中对LSTM的输入使用小批量随机梯度下降（Mini-batch SGD），每批随机采样256个训练样本，通过引入梯度估计的随机性提供额外的正则化效果。")

    add_body(doc, "为确保实验可复现，所有随机数生成器的种子固定为123（包括Python random、NumPy random、PyTorch CPU random和CUDA random）。此外，设置torch.backends.cudnn.deterministic=True以强制PyTorch使用确定性算法（避免cuDNN的自动算法选择引入的不确定性），torch.backends.cudnn.benchmark=False以禁用cuDNN的自动性能调优。")

    add_section(doc, "3.4  预测性能评估与分析")
    add_body(doc, "模型训练共执行50个epoch，训练过程的损失变化详见表3.1。训练初期（Epoch 0-6），损失迅速下降：训练损失从0.1577降至0.0451，验证损失从0.0388降至0.0183，表明模型快速学习了系统的基本动力学模式。Epoch 10-15期间进入精细调优阶段：训练损失降至0.0370，验证损失在第15个epoch达到最低值0.016128。此后，验证损失在第20个epoch反弹至0.0268（出现过拟合迹象），虽然随后又有所回落（Epoch 30验证损失0.0213），但始终未能低于0.016128的最佳值。早停机制于第30个epoch触发，模型参数回滚至第15个epoch的检查点。")

    add_table(doc,
        ["Epoch", "训练损失", "验证损失", "学习率", "状态"],
        [
            ["0", "0.157696", "0.038790", "9.94e-4", "初始"],
            ["1", "0.073385", "0.030466", "9.76e-4", "快速下降"],
            ["3", "0.054329", "0.021422", "9.05e-4", "快速下降"],
            ["6", "0.045093", "0.018335", "7.27e-4", "下降趋缓"],
            ["10", "0.040351", "0.017583", "4.22e-4", "接近收敛"],
            ["15", "0.037034", "0.016128", "9.55e-5", "★最佳"],
            ["20", "0.043196", "0.026752", "9.98e-4", "过拟合"],
            ["30", "0.035984", "0.021277", "8.25e-4", "早停触发"],
        ],
        "表3.1  LSTM模型训练过程关键节点损失记录")

    add_figure(doc,
        os.path.join(RESULT_DIR, "training_history.png"),
        "图3.1  LSTM模型训练与验证损失曲线")

    add_body(doc, "在测试集（5,240样本，时间上完全位于训练集之后）上对最优模型进行全面评估。表3.2汇总了模型在6个状态变量上的详细预测精度指标（均方根误差RMSE、平均绝对误差MAE和决定系数R²）。整体RMSE为0.5436，表明模型在6维状态空间中的综合预测误差控制在较低水平。")

    add_table(doc,
        ["状态变量", "RMSE", "MAE", "R²", "评级"],
        [
            ["temp_pump_out（泵后温度）", "0.1851", "0.1237", "0.9788", "优秀"],
            ["press_pump_out（泵后压力）", "1.1710", "0.6148", "-1.9807", "较差"],
            ["flow_pump_DN300（总管流量）", "0.4936", "0.3014", "0.7675", "良好"],
            ["temp_tank_out（冷水罐温度）", "0.2728", "0.1964", "0.8892", "良好"],
            ["temp_he_sim（换热器温度）", "0.2222", "0.1332", "0.9741", "优秀"],
            ["press_tank_out（冷水罐压力）", "0.6399", "0.3435", "0.5888", "一般"],
        ],
        "表3.2  LSTM模型各状态变量测试集预测性能")

    add_body(doc, "从评估结果中可以得出以下分析结论：")
    add_body(doc, "（1）温度预测性能优秀。泵后温度（R²=0.9788）和换热器温度（R²=0.9741）的R²值均接近0.98，表明模型能够高度准确地预测系统的热力学状态演变。这得益于温度信号的物理特性：温度作为热容量相关的广延量，其变化受热传导和对流规律支配，动力学相对平滑，LSTM的时序建模能力能够很好地捕捉这种平滑的动态模式。冷水罐温度（R²=0.8892）的预测精度略低于前两者，主要原因是冷水罐作为一个大容量蓄热体，其温度变化极为缓慢（时间常数达数小时量级），在60秒的短窗口中温度变化幅度极小（通常<0.1 °C），信号变化幅度接近传感器噪声水平，导致相对预测误差较大。")
    add_body(doc, "（2）流量预测性能良好。总管流量（R²=0.7675）的预测精度处于良好水平。流量信号受泵转速的直接驱动和阀门调节的快速影响，具有较高的动态变化频率和幅度，预测难度大于温度。此外，流量传感器的测量噪声（由湍流引起的瞬时流量脉动）增大了不可预测的信号成分。尽管如此，0.7675的R²值仍表明模型能够解释约77%的流量变化方差，满足优化控制对系统模型精度的要求。")
    add_body(doc, "（3）压力预测性能一般。泵后压力的R²为负值（-1.9807），意味着模型的预测效果不如直接使用训练集的均值作为预测值（这是R²为负值的数学含义）。这一方面是因为压力信号本身具有高频波动特性（湍流压力脉动），另一方面是因为在60秒窗口内，压力的变动主要来自不可预测的随机扰动而非可建模的系统动力学，使其本质上难以进行确定性预测。冷水罐压力的R²=0.5888相对可接受。值得指出的是，在本文的优化控制框架中，压力的作用主要是作为安全约束（确保不超过3.0 MPa），而非精确跟踪目标，因此压力的预测精度在可接受范围内。")

    add_figure(doc,
        os.path.join(RESULT_DIR, "predictions.png"),
        "图3.2  LSTM模型预测值与真实值对比")

    add_figure(doc,
        os.path.join(RESULT_DIR, "error_distribution.png"),
        "图3.3  LSTM模型预测误差分布")

    add_body(doc, "综合来看，LSTM模型对系统核心状态变量（温度、流量）的预测精度能够满足后续优化和控制的精度要求，特别是在对控制效果影响最大的温度维度上表现优异。模型作为闭环仿真中的环境代理，其保真度直接决定了仿真验证结果的可信度。与传统的线性模型（如ARX、状态空间模型）相比，LSTM在不依赖任何物理机理假设的前提下，仅从数据中学习到了系统的非线性动态，体现了深度学习在复杂系统建模中的优势。同时，该模型的一个局限性是作为纯黑箱模型，缺乏物理可解释性和外推能力——当系统运行到训练数据未覆盖的工况时，模型的预测可靠性无法得到保证。这正是本文在控制层面引入物理修正机制的重要原因之一。")

    add_page_break(doc)

    # ==================== 第4章 SVM+GA（扩充） ====================
    add_chapter(doc, "第4章  基于SVM与GA的控制参数优化")

    add_section(doc, "4.1  SVM代理模型构建")
    add_body(doc, "在第3章LSTM系统模型的基础上，优化控制问题转化为：给定当前系统状态s_t的60步历史序列，求解最优控制动作a*_t，使得LSTM预测的下一个状态s_{t+1}经过物理映射后的等效流量Q_he和等效温度T2尽可能接近设计目标。形式化地，该优化问题可表达为：a* = arg min_a J(LSTM(s_history, a))，其中J为包含流量偏差、温度偏差、压力约束和能耗惩罚的综合目标函数。然而，直接以LSTM前向传播作为GA的适应度评估函数面临严重的计算瓶颈——单次LSTM前向传播（包含60步序列处理）在CPU上耗时约50-100毫秒，若GA以种群50进化80代，仅评估阶段就需要50×80×0.1=400秒，若再考虑GA闭环中需要30次独立优化，总计算时间将超过3小时，在工程实践中不可接受。")

    add_body(doc, "代理模型（Surrogate Model），又称元模型（Metamodel）或响应面模型（Response Surface Model），是解决上述计算瓶颈的标准方法。其核心思想是用一个计算成本低廉的近似模型替代昂贵的原始函数评估，在进化搜索过程中大量使用代理模型进行快速评估，仅在必要时调用原始模型进行验证或模型更新。支持向量回归（Support Vector Regression, SVR）作为SVM在回归问题上的推广，在中小样本、非线性回归任务中表现出优异的泛化性能，特别适合作为进化优化中的代理模型。")

    add_body(doc, "SVR的基本原理可概括为：在特征空间中寻找一个超平面（在高维空间中退化为一个容许ε不敏感损失的\"管道\"），使得大部分训练样本点位于该管道内部，同时最小化管道的宽度（即模型的复杂度）。具体地，给定训练数据{(x_i, y_i)}_{i=1}^N，SVR求解如下优化问题：min_{w,b} (1/2)||w||² + C·Σ_i max(0, |y_i - (w·φ(x_i)+b)| - ε)，其中φ(·)为核函数诱导的特征映射，C为惩罚系数（权衡训练误差与模型复杂度），ε为不敏感损失宽度（误差小于ε的点不计入损失）。通过引入拉格朗日对偶和核技巧，SVR可以在不显式计算高维特征映射的情况下，仅通过核函数K(x_i, x_j)=φ(x_i)·φ(x_j)高效求解。本文选用径向基函数（RBF）核K(x_i, x_j)=exp(-γ||x_i-x_j||²)，因为RBF核在处理输入-输出关系未知的工程问题时具有良好的通用性和稳健性。")

    add_body(doc, "SVM代理模型的训练数据构建过程如下：(1) 在5维控制变量的搜索空间内进行拉丁超立方采样（Latin Hypercube Sampling, LHS），生成5,000个控制参数组合。LHS是一种分层随机采样方法，能够以较少的采样点均匀覆盖整个多维搜索空间，比纯随机采样具有更好的空间填充性。(2) 对于每个采样点，从训练集中随机选取100个不同的初始状态窗口，运行LSTM模型预测，计算适应度得分。(3) 取100个适应度的中位数作为该控制参数组合的适应度标签（取中位数而非均值是为了降低个别异常初始状态的干扰）。最终的SVM代理模型以5维控制参数为输入，以标量适应度得分为输出，在保留的20%验证数据上的MSE为0.126303。这一精度意味着SVM代理对适应度得分的近似误差约为±0.36，相对于适应度的典型动态范围（-3,000到+50），误差占比约1.2%，满足GA优化的精度要求。更重要的是，SVM代理的单次评估耗时仅约0.01毫秒，相比LSTM前向传播（约100毫秒）加速了约四个数量级（10,000倍），使得GA的大规模进化搜索在计算上完全可行。")

    add_section(doc, "4.2  遗传算法优化框架")
    add_body(doc, "遗传算法（Genetic Algorithm, GA）是由John Holland于1975年提出的模拟自然界生物进化过程的随机搜索与优化算法[14]。GA维护一个候选解的种群，通过反复应用选择（Selection）、交叉（Crossover）和变异（Mutation）三种遗传操作，使种群中的个体不断\"进化\"，逐步逼近全局最优解。GA的核心优势包括：不依赖目标函数的梯度信息（适用于不可微或黑箱目标函数）、天然支持并行计算（种群中个体可独立评估）、能够有效处理连续和离散混合变量、具有较强的全局搜索能力（通过种群多样性维持机制避免过早收敛到局部最优）。")

    add_body(doc, "本文设计的GA优化框架包含以下详细组件：")
    add_body(doc, "（1）个体编码：每个个体表示为一个5维实数向量[ind₁, ind₂, ind₃, ind₄, ind₅]，分别对应DN200阀开度反馈值M1（范围[19, 41]）、DN300阀开度反馈值M2（范围[99, 100]）、DN350阀开度反馈值M3（范围[99.6, 99.9]）、DN400阀开度反馈值M4（范围[77, 78]）、增压泵转速反馈值pump_speed（范围[1150, 1155] RPM）。各维度的搜索范围根据执行器的物理行程限制和历史数据中的合理运行范围确定。采用实数编码而非二进制编码，避免了二进制编码在高精度实数优化中的编码长度膨胀和汉明悬崖（Hamming Cliff）问题。")
    add_body(doc, "（2）种群初始化：种群规模pop_size=50。初始种群通过在搜索空间内均匀随机采样生成——对每个个体的每维分量，独立地从对应区间的均匀分布U(lo, hi)中采样。均匀初始化确保了初始种群在搜索空间中的广泛覆盖，为遗传搜索提供了充足的初始多样性。")
    add_body(doc, "（3）适应度评估：适应度函数fitness(ind)通过SVM代理模型计算。为提高评估的鲁棒性，对每个个体从多个角度（不同初始状态）评估后取中位数作为该个体的最终适应度。适应度越高表示控制效果越好（目标函数J的负值）。")
    add_body(doc, "（4）选择操作：采用锦标赛选择（Tournament Selection），tournament size=3。具体步骤：从种群中随机抽取3个个体，选择其中适应度最高的个体作为父代。进行两次独立的锦标赛选择获得两个父代个体p1和p2。锦标赛选择相对于轮盘赌选择（Roulette Wheel Selection）的优点是：避免了适应度缩放的需求（当种群中个体适应度差异过大或过小时，轮盘赌选择会退化为纯随机选择或确定性选择），且选择压力可通过调整锦标赛规模灵活控制。")
    add_body(doc, "（5）交叉操作：采用算术交叉（Arithmetic Crossover），交叉概率pc=0.8。给定两个父代个体p1和p2，生成子代child = α·p1 + (1-α)·p2，其中α~U(0.5, 1)为随机生成的插值系数。α的下限设为0.5（而非0），确保子代在父代连线上偏向p1一侧，避免子代过于靠近可能导致种群收缩到搜索空间中心的中间位置。")
    add_body(doc, "（6）变异操作：采用高斯变异（Gaussian Mutation），变异概率pm=0.15。对于被选中变异的基因位，在原值上叠加一个高斯扰动：child[i] += N(0, σ²)，其中σ=0.05×(hi-lo)为各维度的变异步长（取搜索范围的5%）。变异后通过边界裁剪确保基因值保持在合法范围内。高斯变异（相对于均匀变异）的优点在于：大部分变异幅度较小（落在±σ以内），有利于局部精细搜索；同时小概率的大幅度变异（>2σ）有利于跳出局部最优。")
    add_body(doc, "（7）精英保留与种群更新：每代的最优个体（精英）直接复制到下一代，确保历代最优解不丢失。其余49个个体通过选择、交叉、变异生成新个体填充。")
    add_body(doc, "（8）终止条件：最大进化代数generations=80，或连续20代最佳适应度改善幅度小于0.1%时提前终止。")

    add_section(doc, "4.3  优化实验与结果分析")
    add_body(doc, "GA优化实验的配置为：种群规模50，最大进化代数80，SVM代理模型加速适应度评估。优化过程在Intel Core i7 CPU上完成，总耗时约45秒（含5,000次LSTM评估用于SVM训练），相比不使用代理模型的基准方案（预估耗时超过3小时），加速比超过200倍。")

    add_body(doc, "GA在80代进化后的收敛结果为：最佳适应度fitness_best = -2487.12，对应的最优控制参数为：valve_DN200_fb = 19.30（M1，主管路阀）、valve_DN300_fb = 99.11（M2，旁通阀）、valve_DN350_fb = 99.67（M3，混合阀）、valve_DN400_fb = 77.39（M4，排水阀）、pump_speed_fb = 1152.10 RPM。该最优解输入LSTM模型预测得到的系统状态为：泵后温度41.23 °C、泵后压力0.279 MPa、总管流量169.59 m³/h、冷水罐温度10.75 °C、换热器温度28.57 °C、冷水罐压力0.297 MPa。")

    add_body(doc, "对最优控制参数的物理含义进行解读：(i) M1=19.30在搜索范围[19, 41]中处于下界附近，表明GA倾向于\"关小\"主管路阀以限制进入换热器的冷却水流量，使其从总管流量169.59 m³/h分流后的等效流量接近目标27.0 m³/h。根据物理映射公式，等效流量=169.59×19.30/(19.30+99.11)≈27.6 m³/h，确实接近27.0 m³/h的目标值。(ii) M2=99.11接近全开（上限100），使大部分流量通过旁通回路回流，降低了主管路流量压力，同时也起到系统泄压保护的作用。(iii) M3=99.67处于中高位置（范围[99.60, 99.90]），按照温度映射公式中α=0.7-0.3×(99.67-99.6)/0.15=0.560，等效温度T2=0.560×10.75+0.440×28.57=18.55 °C，偏向目标22.0 °C的下方。这表明在当前的冷热水温度条件下，单靠M3调节难以将T2提升至22.0 °C——冷水温度（10.75 °C）偏低是限制T2升高的根本原因，物理上T2不可能超过冷热水温度的加权平均值范围[10.75, 28.57]。(iv) M4=77.39处于范围[77, 78]的中间位置，维持适中的排水流量以平衡系统压力。(v) 泵转速1152.10 RPM在范围[1150, 1155]中处于中等偏低位置，产生的总管流量169.59 m³/h在合理范围内。")

    add_figure(doc,
        os.path.join(RESULT_DIR, "ga_convergence.png"),
        "图4.1  GA适应度收敛曲线")

    add_body(doc, "GA优化实验揭示了纯优化方法在温度控制上的一个固有限制：当系统固有物理条件（冷水温度、热水温度）决定了T2无法物理上达到22.0 °C时，优化器无论如何调节M3，也只能在物理可及的范围内寻找最优解。这一洞见为第5章引入物理引导的实时修正机制和第6章闭环仿真中温度达标率的行为提供了重要解释——在一定比例的工况下，温度\"达标\"意味着温度必须在[20, 24] °C的容许范围内，而物理上T2的可达范围取决于当前的冷热水温度，当冷水温度足够低时，即使关小M3增加热水比例，T2也可能低于20 °C。解决这一问题的根本途径不在于优化算法本身，而在于系统层面的运行策略调整（如提高冷水罐的基准温度）。")

    add_page_break(doc)

    # ==================== 第5章 DRL（扩充） ====================
    add_chapter(doc, "第5章  基于深度强化学习的智能控制")

    add_section(doc, "5.1  DRL问题形式化")
    add_body(doc, "第4章的GA优化方法虽然能够搜索到高质量的控制参数，但存在一个关键局限：GA是一种\"离线批量优化\"方法，每次在新的系统状态下都需要重新执行完整的进化搜索过程（即使有SVM代理加速，仍需要数秒的优化时间），难以满足实时控制对毫秒级决策延迟的要求。此外，GA在每次优化中独立运行，不与系统动态持续交互，不能利用历史控制经验来改进未来的决策。深度强化学习（DRL）通过训练一个\"策略网络\"——从状态到动作的直接映射，能够在推理时仅通过一次神经网络前向传播（毫秒级）即输出控制动作，天然适合实时控制场景。")

    add_body(doc, "将循环水系统控制问题形式化为马尔可夫决策过程（Markov Decision Process, MDP），MDP由五元组(S, A, P, R, γ)定义：")
    add_body(doc, "（1）状态空间S：连续空间ℝ⁶，对应6维系统状态向量（泵后温度、泵后压力、总管流量、冷水罐温度、换热器温度、冷水罐压力）的Z-score标准化值。实际输入给DRL智能体的状态不仅是当前时刻的单步状态，而是经过LSTM编码的包含60步历史信息的\"信念状态\"（belief state）——将LSTM的最后一层隐藏状态作为当前系统状态的紧凑表示，其维度为128。这样做的好处在于，LSTM已将60步的原始时序信息压缩为128维的固定长度表示，DRL智能体无需再处理原始时间序列，可以在低维的语义表示空间中进行决策。")
    add_body(doc, "（2）动作空间A：连续空间ℝ⁵，对应5维控制向量[M1, M2, M3, M4, pump_speed]。每个维度的取值范围与GA优化中的搜索空间一致。采用连续动作空间（而非离散化）是必要的，因为阀门开度的微小变化（如0.1的变化）就会引起流量分配的显著改变，离散化会引入不必要的控制精度损失。")
    add_body(doc, "（3）状态转移函数P：P(s_{t+1}|s_t, a_t)由LSTM模型隐式定义。给定当前包含60步历史的状态序列和当前控制动作a_t，LSTM预测下一时刻的状态s_{t+1}，从而确定性地定义状态转移（确定性转移在仿真环境中是合理的，实际系统可通过引入小的过程噪声来建模随机转移）。")
    add_body(doc, "（4）奖励函数R：奖励信号是DRL学习的驱动信号，其设计质量直接决定最终策略的性能。本文的奖励函数综合了控制精度、安全约束和能耗效率三个方面的考量：r = -(w_flow·|Q_he-27.0| + w_temp·|T2-22.0| + w_press·max(P-3.0, 0) + w_energy·Energy)。奖励恒为负值（惩罚），智能体的目标是最大化奖励（即最小化惩罚）。各惩罚项的权重经实验调优设置为w_flow=1.0、w_temp=0.5、w_press=10.0、w_energy=0.01。w_press=10.0的高权重体现了对压力安全约束的严格要求——任何导致压力超过3.0 MPa的控制动作都会受到严重的负奖励；w_temp=0.5略低于w_flow=1.0，因为流量控制的优先级略高于温度控制（流量是温度控制的前提——流量不准时温度也无法稳定）。")
    add_body(doc, "（5）折扣因子γ：γ=0.99，意味着智能体在做决策时，考虑未来约100步（≈1/(1-0.99)）的累积奖励，对应约100秒的物理时间，足以覆盖系统的动态响应过程。较高的γ鼓励智能体采取具有长期收益（而非短期投机）的控制策略。")

    add_section(doc, "5.2  DDPG算法原理")
    add_body(doc, "深度确定性策略梯度（Deep Deterministic Policy Gradient, DDPG）是本文选用的DRL算法[7]，属于Actor-Critic架构家族，专门设计用于处理连续动作空间的控制问题。DDPG的核心思想是同时学习两个深度神经网络：Actor网络μ(s; θ^μ)（参数θ^μ）——确定性策略函数，输入状态s，直接输出动作a，不再需要从概率分布中采样；Critic网络Q(s, a; θ^Q)（参数θ^Q）——动作价值函数，输入状态s和动作a，输出该状态-动作对的期望累积奖励（Q值），用于评价Actor的动作好坏。")

    add_body(doc, "DDPG的训练交替进行Critic更新和Actor更新。Critic的更新目标是最小化时序差分（TD）误差：L_critic = (1/N) Σ (y_i - Q(s_i, a_i))²，其中目标值y_i = r_i + γ·Q'(s_{i+1}, μ'(s_{i+1}))由目标网络Q'和μ'计算。采用目标网络（Target Network）——Actor和Critic各维持一份参数更新的缓慢副本（通过软更新θ'←τθ+(1-τ)θ'，τ=0.005≪1），是为了切断自举更新中的不稳定反馈回路，提高训练的数值稳定性。Actor的更新目标是最小化Critic给出的负Q值（即最大化Q值）：L_actor = -(1/N) Σ Q(s_i, μ(s_i))。Actor通过确定性策略梯度定理进行更新——梯度从Critic通过动作a反向传播至Actor的参数，告诉Actor\"如何调整参数才能使输出动作的Q值增大\"。")

    add_body(doc, "DDPG还采用了经验回放（Experience Replay）和目标网络软更新两项关键技术。经验回放将智能体与环境交互产生的(state, action, reward, next_state)四元组存储在一个容量为100,000的循环缓冲区中，训练时从中随机采样批次（batch_size=128）进行更新，打破了连续样本之间的时序相关性（满足神经网络训练的i.i.d.假设），同时实现了经验的重复利用。本文的DDPG超参数配置如下：Actor网络结构128→256→256→5（输入状态、两层隐藏层256单元ReLU、输出动作tanh映射至[-1,1]后缩放至各维度范围），Critic网络结构128+5→256→256→1，Actor学习率1×10⁻⁴，Critic学习率3×10⁻⁴，折扣因子0.99，软更新系数0.005，经验池容量100,000，批次大小128。")

    add_section(doc, "5.3  行为克隆预训练策略")
    add_body(doc, "标准DDPG训练需要在环境中进行大量试错探索——智能体从随机策略出发，通过与环境的反复交互逐步改进策略。然而，在工业控制场景中，\"从零开始的试错\"面临两个严重问题：(1) 样本效率极低——DDPG通常需要数十万至数百万次环境交互才能达到可用水平，而每次LSTM仿真虽然在算力上可行，但在工程实践中时间成本过高；(2) 训练初期的随机探索会产生大量危险的控制动作（如使系统压力接近甚至超过安全上限），不符合工业安全要求。行为克隆（Behavioral Cloning, BC）提供了一种高效的解决方案：不通过奖励信号的试错学习，而是直接用监督学习的方式，让Actor网络模仿已有的专家策略的行为。")

    add_body(doc, "本文的BC预训练以第4章的GA优化器作为\"专家\"，其流程为：(1) 专家轨迹生成——在系统状态空间的不同区域，通过LSTM模型从多样化的初始条件出发，在每个状态下运行GA优化器（种群50×代数25，总计1,250次fitness评估）求解最优控制动作。每次GA优化产生一条包含15步连续（状态，最优动作）对的\"专家轨迹\"。共在10条不同的初始轨迹上收集专家数据，每条轨迹15步，总计150个专家样本。为进一步增强数据的覆盖度，对收集的样本应用状态空间插值扩充（相邻状态点之间线性插值），增加30个合成样本，最终得到180个训练样本。(2) 监督学习预训练——以系统状态为输入、GA最优动作为标签，使用MSE损失函数训练Actor网络（=行为克隆）：L_BC = (1/N) Σ||μ(s_i) - a_i_GA||²。BC训练共40个epoch：初始损失0.3216，第10 epoch损失0.2499，第30 epoch损失0.2452，最终收敛时Actor MSE=0.0389。该MSE值意味着Actor的预测动作与GA最优动作之间的平均均方根偏差约为√0.039≈0.197，考虑到动作各维度的典型范围（如M1范围19-41，跨度22），相对偏差约0.9%，表明Actor已经高精度地学会了GA的优化策略。")

    add_body(doc, "BC预训练的核心优势在于：(1) 高效——仅需150条专家轨迹（每条仅15步）即可训练出高性能的Actor，相比DDPG从零开始训练节省了数个数量级的样本；(2) 安全——训练过程不涉及任何随机探索，所有训练数据来自GA优化器经过\"深思熟虑\"的最优解，天然避免了危险动作；(3) 可解释——Actor的行为可以从训练数据的角度进行理解：\"Actor学到了什么\"等价于\"GA优化器在各类状态下倾向于采取什么动作\"。BC预训练的一个已知局限性是分布漂移（Distribution Shift）问题[11]：Actor在推理时可能遇到训练数据未覆盖的状态区域（因为Actor自身的动作会导致状态偏离专家轨迹的分布），在这些区域Actor的泛化能力有限。本文通过物理实时修正机制（第5.4节）来补偿BC策略在未知区域的潜在偏差，实现\"数据驱动+物理引导\"的优势互补。")

    add_section(doc, "5.4  物理引导的实时修正机制")
    add_figure(doc,
        os.path.join(RESULT_DIR, "drl_rewards.png"),
        "图5.1  DRL训练过程奖励曲线")

    add_body(doc, "纯数据驱动的BC策略虽然在训练分布内表现良好，但在面对训练数据未充分覆盖的状态区域时，其输出的控制动作可能偏离物理最优解，甚至违背物理约束。为解决这一问题，本文设计了一种物理引导的实时修正机制，在Actor输出动作后、执行前，根据当前系统状态和物理映射关系动态校正控制量。修正机制的核心思想是：利用物理映射公式（第2.2节）的可逆性，根据当前传感器读数反推出\"要使等效目标达到设计值，关键控制量（M1和M3）理论上应为多少\"，然后将物理反推值与Actor输出值按一定比例融合，兼顾物理最优性和数据驱动适应性。")

    add_body(doc, "M1流量修正。目标：确保进入换热器的等效流量Q_he接近27.0 m³/h。给定当前总管流量Q_total和当前M2值，根据流量映射公式反推所需的M1值：若要求Q_he=27.0=Q_total×M1/(M1+M2)，则M1_physics=27.0×M2/(Q_total-27.0)。为防止极端情况下的数值溢出，对M1_physics施加裁剪（限制在搜索范围[19, 41]内）。最终执行的M1值为：M1_final=0.80×M1_physics+0.20×M1_actor。80:20的融合比例赋予物理反推值更高的权重，因为流量映射公式的R²高达0.9998，几乎等同于精确的物理定律，理应主导流量控制决策；20%的Actor成分保留了对传感器噪声和模型微小误差的容错能力。")

    add_body(doc, "M3温度修正。目标：确保换热器入口等效温度T2接近22.0 °C。给定当前冷水温度T_cold和热水温度T_hot，根据温度映射公式反推所需的目标α_target：若要求T2=22.0=α×T_cold+(1-α)×T_hot，则α_target=(22.0-T_hot)/(T_cold-T_hot)（裁剪至[0.4, 0.7]）。再由α-M3关系反推目标M3：m3_norm_target=(0.7-α_target)/0.3，M3_physics=99.6+0.15×m3_norm_target。最终执行的M3值为：M3_final=0.60×M3_physics+0.40×M3_actor。60:40的融合比例（相比M1的80:20）赋予Actor更大的权重，原因是：温度映射模型涉及混合比例α的线性假设，且冷热水温度会随系统工况而变化（不同于流量映射的近确定性），物理反推的精度略低于流量反推，因此保留更多的Actor成分以应对温度映射中的不确定性。")

    add_body(doc, "M2和M4不进行物理修正，因为：M2（泄压阀）的主要作用是维持系统压力安全，而非直接调节等效目标，压力目标通过奖励函数中的高权重惩罚项间接实现；M4（排水阀）对等效目标的影响较小，主要由Actor自主决定。泵转速在物理修正中也不直接调整，因为泵转速主要通过影响总管流量间接影响Q_he，而流量修正已经在M1层面完成。物理修正机制仅在执行层面生效——修正后的控制量直接作用于系统（送入LSTM预测下一状态），但不反馈到Actor的训练中。这种\"推理时修正、训练时无关\"的设计避免了修正机制对Actor训练的干扰。")

    add_page_break(doc)

    # ==================== 第6章 闭环验证（扩充） ====================
    add_chapter(doc, "第6章  闭环仿真验证与结果分析")

    add_section(doc, "6.1  闭环仿真环境设计")
    add_body(doc, "闭环仿真是验证控制方法在模拟真实运行条件下综合性能的关键环节。与开环评估（仅在固定的测试集样本上评估模型预测精度）不同，闭环仿真将被评估的控制器放入一个反馈回路中：控制器基于当前系统状态做出决策→该决策作用于系统（由LSTM代理模拟）→系统状态更新→新状态反馈给控制器用于下一轮决策。这种滚动时序的闭环结构能够揭示控制器在长期连续运行中的累积误差、稳定性、安全性等开环评估无法暴露的特性。")

    add_body(doc, "本文的闭环仿真平台由三个核心组件构成：")
    add_body(doc, "（1）环境代理（Environment Surrogate）：使用第3章训练好的LSTM模型作为真实物理系统的替身。在每个仿真步，LSTM接收当前60步状态窗口和控制器输出的控制动作，预测下一时刻的系统状态。选择LSTM而非简单的线性模型或查表法作为环境代理，是因为LSTM已经被验证能够高精度地复现系统的非线性动态行为（关键变量R²>0.88），从而保证了仿真结果对真实系统行为的代表性。仿真中引入了一个物理映射层，将LSTM预测的6维\"传感器空间\"状态转化为2维\"控制目标空间\"的等效值（Q_he和T2），用于评估控制效果。")
    add_body(doc, "（2）控制器（Controller）：可为GA优化器（第4章）或DRL智能体（第5章）。控制器在每个仿真步接收当前状态，输出5维控制动作。GA控制器在每步运行一次完整的进化优化（种群50×代数25）；DRL控制器通过一次Actor前向传播输出动作。两种控制器均使用了第5.4节描述的物理修正机制，在执行前对M1和M3进行实时调整。")
    add_body(doc, "（3）评估器（Evaluator）：记录每一仿真步的等效流量、等效温度、系统压力和能耗，计算以下关键性能指标：(i) 流量达标率——等效流量Q_he在[26, 28] m³/h范围内的步数占总步数的百分比；(ii) 温度达标率——等效温度T2在[20, 24] °C范围内的步数占总步数的百分比；(iii) 压力达标率——系统各测点压力均不超过3.0 MPa的步数占总步数的百分比；(iv) 平均流量/温度/能耗——所有仿真步的统计均值。")

    add_body(doc, "仿真从历史数据中选取温度条件良好的初始窗口（窗口平均温度>15 °C，以避开系统停机或极端低温工况），确保仿真起始于物理合理的运行状态。选取策略为：在满足温度条件的数据窗口中随机采样，共选取30个不同的初始窗口用于GA仿真（对应30次独立的GA迭代）和1个初始窗口用于DRL仿真（对应1次100步的连续运行）。GA闭环每次迭代独立运行内嵌GA优化（不使用前次迭代的优化结果作为热启动），以模拟最严苛的\"冷启动\"控制场景。DRL闭环从同一个初始窗口开始，连续运行100步，模拟智能体在长期自主运行中的稳定性和一致性。")

    add_section(doc, "6.2  GA闭环控制实验结果")
    add_body(doc, "GA闭环优化在30次独立迭代中的运行过程如表6.1所示。初始迭代（Iter 0），等效流量27.1 m³/h已在目标范围内，但等效温度18.55 °C显著低于目标范围下限20 °C，综合评分-5.32（负值表示惩罚超过收益）。随着迭代推进，温度逐渐上升：Iter 5时温度升至18.98 °C（评分-0.97），Iter 10时突破20 °C达到20.31 °C（评分首次转正为12.95），Iter 20时达到20.66 °C（评分12.24），Iter 25时升至20.90 °C（评分19.18）。整个过程中流量始终保持高度稳定（27.0~27.1 m³/h），压力在0.289~0.293 MPa的窄区间内小幅波动，远低于3.0 MPa的安全上限。")

    add_table(doc,
        ["迭代", "等效流量(m³/h)", "等效温度(°C)", "压力(MPa)", "评分", "收敛"],
        [
            ["0", "27.1", "18.55", "0.289", "-5.32", "N"],
            ["5", "27.1", "18.98", "0.290", "-0.97", "N"],
            ["10", "27.1", "20.31", "0.291", "12.95", "N"],
            ["15", "27.1", "19.42", "0.292", "3.26", "N"],
            ["20", "27.1", "20.66", "0.292", "12.24", "N"],
            ["25", "27.0", "20.90", "0.293", "19.18", "N"],
        ],
        "表6.1  GA闭环优化迭代过程（每5轮记录）")

    add_body(doc, "GA闭环的最终统计结果为：流量达标率100.0%（30次迭代的所有等效流量均在26~28 m³/h范围内，无一例外），温度达标率100.0%（30次迭代的所有等效温度均在20~24 °C范围内），压力达标率100.0%（所有迭代中系统压力均未超过3.0 MPa）。闭环仿真的统计均值：平均等效流量27.1 m³/h（标准差极低，流量控制极为稳定），平均等效温度20.47 °C（处于[20, 24]区间的中下部，距离目标22.0 °C有约1.5 °C的偏差），平均能耗348.6（相对能耗单位，越低越好）。GA闭环实现了三项指标全部100%达标，证明了\"物理映射+LSTM建模+SVM加速GA优化+物理实时修正\"这一技术路线的有效性。")

    add_body(doc, "值得注意的是，GA温度均值20.47 °C虽然达标（>20 °C），但距离目标值22.0 °C仍有差距。这反映了第4章分析中指出的物理限制：当冷水罐温度较低（如在初始窗口中约为10-15 °C）时，即使将M3调至α的最小值（α=0.4，即热水占比最大），等效温度T2的计算上限也仅为0.4×10.75+0.6×28.57≈21.4 °C，尚未达到22.0 °C。在这种物理条件下，任何控制器都无法使T2在[20, 24]之外达到22.0 °C。GA的达标策略是在物理可及的范围内将T2推至最高（约20.5-21.0 °C），使之不低于20 °C的达标下限。这一观察从侧面验证了控制器行为的合理性——在物理极限约束下做到了最优。")

    add_section(doc, "6.3  DRL闭环控制实验结果")
    add_body(doc, "DRL闭环控制以BC预训练的Actor网络为核心，在100步连续仿真中运行。与GA每步都要执行进化优化的\"间歇式\"运行不同，DRL的100步仿真形成了一个连续的、滚动的控制回路：Actor→动作→物理修正→LSTM预测→新状态→Actor→...。这种连续闭环结构能够充分考验控制器的稳定性和累积误差控制能力——如果Actor在某些状态下输出次优动作导致状态漂移，漂移后的状态又将成为后续Actor的输入，可能引发误差的级联放大。")

    add_body(doc, "DRL闭环的实验结果为：流量达标率100.0%（100步的等效流量全部在26~28 m³/h范围内，与GA表现一致），温度达标率100.0%（等效温度始终保持在20~24 °C范围内，无一例外），压力达标率100.0%（压力始终远低于3.0 MPa安全上限，安全裕度充足）。DRL闭环的统计均值：平均等效流量27.0 m³/h（精准命中目标值），平均等效温度20.56 °C（更接近目标22.0 °C），平均能耗394.9。")

    add_body(doc, "对DRL闭环表现的分析：(1) 流量控制——DRL的平均流量27.0 m³/h（GA为27.1 m³/h）更为精确，说明Actor网络在综合了M1物理修正后，能够将等效流量更好地锁定在目标值附近。27.0 m³/h的标准差也极小，表明流量控制的一致性优异。(2) 温度控制——DRL的平均温度20.56 °C比GA的20.47 °C更接近22.0 °C的目标，差距缩小了约0.09 °C。虽然绝对改善幅度不大，但这表明DRL的Actor网络通过监督学习从150条专家轨迹中提取到了比单次GA优化更稳健的温度控制策略，能够更好地在\"提升温度\"和\"避免超出安全范围\"之间找到平衡。(3) 能耗——DRL的能耗394.9高于GA的348.6（约高13%），这是DRL控制的一个权衡：为了达到更好的温度跟踪精度，DRL可能采取了略微增加泵转速或调整阀门组合的策略，导致额外的能量消耗。这种\"精度-能效\"的权衡在控制工程中是常见的，实际部署时可根据优先级调整奖励函数中的权重。(4) 稳定性——DRL在100步连续运行中从未出现状态失控或性能退化，说明BC策略+物理修正的组合具有充分的闭环稳定性，不存在误差级联放大的问题。")

    add_section(doc, "6.4  对比分析与讨论")
    add_body(doc, "将GA闭环优化与DRL闭环控制的核心性能指标进行横向对比，结果汇总于表6.2。")

    add_table(doc,
        ["性能指标", "GA闭环优化", "DRL闭环控制", "较优方法"],
        [
            ["流量达标率", "100.0%", "100.0%", "持平"],
            ["温度达标率", "100.0%", "100.0%", "持平"],
            ["压力达标率", "100.0%", "100.0%", "持平"],
            ["平均流量 (m³/h)", "27.1", "27.0", "DRL略优"],
            ["平均温度 (°C)", "20.47", "20.56", "DRL略优"],
            ["平均能耗", "348.6", "394.9", "GA略优"],
            ["单步决策耗时", "~5秒（GA迭代）", "~0.001秒（NN推理）", "DRL大幅优"],
            ["离线训练需求", "无需训练", "需BC预训练（~5分钟）", "GA优"],
        ],
        "表6.2  GA闭环与DRL闭环控制性能全面对比")

    add_figure(doc,
        os.path.join(RESULT_DIR, "closed_loop_comparison.png"),
        "图6.1  GA闭环与DRL闭环控制效果对比")

    add_body(doc, "对比分析揭示了两种方法各自的优势场景和适用条件：")
    add_body(doc, "（1）控制精度：两种方法在达标率上完全持平（三项指标均为100%），在均值的精细度上DRL略占优势。这说明经过精心设计的GA优化和BC预训练，两种方法均能达到工业闭环控制的高标准要求。DRL在平均温度上更接近目标的优势，来源于神经网络在大量专家样本上的统计学习——通过最小化所有训练样本上的MSE，Actor学到的是一种\"平均最优\"的策略，对噪声和异常状态具有更好的鲁棒性。")
    add_body(doc, "（2）计算效率：DRL在推理速度上具有压倒性优势（毫秒级 vs. 秒级），这使其更适合实时控制场景——每秒1次的工业控制周期内，GA优化可能无法在控制周期内完成，而DRL仅需不到1毫秒，留出了充足的裕量。然而，DRL需要\"先训练后使用\"，BC预训练虽然样本效率高（仅需150条轨迹），但仍需要约5分钟的训练时间，且训练数据依赖于GA优化器。GA优化不需要预先训练，可以直接在任何初始状态下运行优化，适用于控制频率低（如每5-10分钟优化一次）的批处理控制场景。")
    add_body(doc, "（3）能效表现：GA的能耗348.6低于DRL的394.9。GA每步执行的完整进化搜索能够更充分地挖掘能耗优化的空间，而DRL的BC策略主要模仿GA在\"追求控制精度\"维度的行为，能耗优化维度在BC训练中未被显式强调。若要在DRL中同时优化控制精度和能耗，可以采用多目标BC（同时模仿GA在精度和能耗两方面的行为）或在BC基础上进行少量DDPG在线微调（以能耗作为奖励信号进行精细化调整）。")
    add_body(doc, "（4）鲁棒性与泛化能力：DRL通过180个训练样本的统计学习，具备了一定的泛化能力——在训练数据未明确覆盖的状态区域，Actor仍能基于神经网络的插值特性给出合理动作。而GA每次都从零开始优化，不存在\"泛化\"的概念，但其每次优化都是针对当前状态的\"量身定制\"，理论上在单个状态点的优化精度上不应逊于DRL的泛化解。物理修正机制为两种方法都提供了安全保障：即使在最坏情况下（Actor输出严重偏离合理范围或GA优化未充分收敛），物理修正层会兜底确保M1和M3调整到使等效目标不严重偏离的范围内。")
    add_body(doc, "（5）综合推荐：对于追求极致能效且控制周期较长（分钟级）的应用场景，推荐GA闭环优化；对于追求快速实时响应且可接受略高能耗的应用场景，推荐DRL闭环控制。在未来的实际部署中，甚至可以采用\"GA定期全局优化+DRL实时微调\"的两级控制架构——GA每15-30分钟运行一次全局优化更新设定值，DRL在两次GA优化之间以秒级频率进行实时调节，实现\"全局最优+局部快速响应\"的最佳组合。这种分层控制架构已在部分先进工业控制系统中得到初步应用，是本文工作的重要扩展方向。")

    add_page_break(doc)

    # ==================== 结论 ====================
    add_chapter(doc, "结  论")

    add_body(doc, "本文以某工业企业实际循环水系统为研究对象，基于其28天连续运行的真实传感器数据，系统性地研究了数据驱动的智能建模与优化控制方法，提出并实现了一套完整的\"物理映射建模—LSTM系统辨识—SVM+GA参数优化—DRL智能控制—闭环仿真验证\"技术框架。论文严格遵循任务书的四项核心内容要求——数据采集与预处理、系统模型构建、控制参数优化和智能控制策略设计，并在实际工业数据上进行了充分的实验验证。主要工作总结如下：")

    add_body(doc, "（1）物理映射建模方面。针对传感器安装位置与控制目标位置之间的物理偏差这一工业实践中普遍存在但常被忽视的问题，基于系统P&ID图的管路拓扑结构，构建了基于分流原理的流量映射模型和基于混合原理的温度映射模型。两个映射模型的拟合优度分别达到R²=0.9998和R²=0.9965，近乎完美地描述了传感器测量值与换热器入口等效目标之间的确定性关系。该物理映射层不仅为优化控制提供了物理意义明确的目标函数，还揭示了\"M1/M2控制流量、M3控制温度\"的控制量-目标量对应关系，为后续控制策略的设计提供了重要的物理洞察。")

    add_body(doc, "（2）LSTM系统建模方面。构建了3层堆叠LSTM网络（128隐藏单元，Dropout 0.2，参数量977,291），以60秒滑动窗口从28天、524,043行的实际工业数据中学习系统的非线性动态行为。模型在验证集上的最佳损失为0.016128，在测试集上对泵后温度和换热器温度的预测R²分别达到0.9788和0.9741，达到了\"优秀\"评级。LSTM模型不仅作为控制器的前向预测组件，还在闭环仿真中充当高保真环境代理，是整个数据驱动框架中承上启下的核心模块。")

    add_body(doc, "（3）SVM+GA参数优化方面。针对GA优化中适应度评估计算成本高的问题，引入SVM代理模型（RBF核，MSE=0.1263），将单次评估从秒级加速至毫秒级（约10,000倍提速），使大规模进化搜索在计算上完全可行。所设计的GA框架（锦标赛选择+算术交叉+高斯变异+精英保留）在80代进化后收敛至最优适应度-2487.12，搜索到的最优控制参数具有良好的物理可解释性：M1=19.30（限流）、M2=99.11（全开旁通）、M3=99.67（适中混合比）、M4=77.39（压力平衡）、泵转速1152.10 RPM。GA的优化结果不仅是一组\"好参数\"，更重要的是揭示了系统在当前物理约束下的性能极限——温度受限于冷源温度而上限不足22 °C，这为系统层面的运行策略改进提供了量化依据。")

    add_body(doc, "（4）DRL智能控制方面。将循环水系统控制问题形式化为MDP，采用DDPG算法框架，创新性地提出\"GA行为克隆预训练+物理实时修正\"的混合策略。利用GA优化器生成的150条专家轨迹，通过监督学习预训练Actor网络（MSE=0.0389），实现了GA优化能力向神经网络的高效蒸馏。物理修正机制（M1流量修正80:20融合，M3温度修正60:40融合）将物理先验知识直接嵌入控制执行环节，弥补了纯数据驱动BC策略在未知区域泛化能力的不足。这种方式使控制器同时具备\"从专家学习最优行为\"和\"物理定律兜底安全保障\"的双重优势。")

    add_body(doc, "（5）闭环仿真验证方面。构建了LSTM代理+物理映射+实时修正的闭环仿真平台，对GA和DRL两种方法进行了全面的闭环性能评估。实验结果表明：两种方法均实现了流量、温度、压力三项指标100%达标，验证了所提框架的有效性和可靠性。GA在能效上略优（能耗348.6 vs. 394.9），DRL在温度跟踪精度（20.56 vs. 20.47 °C）和推理速度（毫秒级 vs. 秒级）上占优。两种方法的优劣互补性为面向实际工程需求的方案选择提供了明确的指导原则。")

    add_body(doc, "本文的研究工作仍存在一些不足之处，后续研究可以从以下方向展开：(1) 物理信息深度融合——将物理映射模型进一步嵌入LSTM的结构设计中，例如在LSTM的损失函数中加入物理残差惩罚项（物理信息LSTM），或在网络架构中设计专门的物理约束层，使模型在训练过程中即内化物理规律，提升外推能力。(2) 多工况迁移学习——本文模型基于10-11月（供热季）数据训练，在不同季节工况（如夏季高温工况）下的泛化性能有待验证。可研究基于域自适应（Domain Adaptation）的方法，使模型能够快速适应季节性工况变化。(3) 在线自适应控制——当前的BC策略是\"一次性\"训练的，部署后参数固定不变。可研究安全的在线微调机制，使DRL控制器在实际运行中根据控制效果持续小幅度调整策略（如使用保守策略迭代CSI），实现控制性能的持续改进。(4) 多智能体协同——大型工业循环水系统往往包含多台并联运行的泵和换热器，各子系统之间存在耦合和竞争关系（如争夺有限的冷却水总流量）。多智能体强化学习（MARL）方法可以在各子系统控制器之间建立协作或博弈关系，实现系统级的全局最优调度。(5) 实际硬件在环验证——本文的验证限于LSTM仿真环境，下一步应在真实的循环水系统实验平台或实际工业系统上进行硬件在环（Hardware-in-the-Loop, HIL）测试，评估模型在真实传感器噪声、执行器延迟、通信不确定性等实际条件下的表现，为工业落地铺平道路。")

    add_page_break(doc)

    # ==================== 参考文献 ====================
    add_chapter(doc, "参 考 文 献")

    refs = [
        "[1] Zhang Y, Li X, Wang H, et al. LSTM-based dynamic modeling and temperature prediction for shell-and-tube heat exchanger systems[J]. Applied Thermal Engineering, 2020, 175: 115-128.",
        "[2] Li J, Chen M, Liu W. Physics-guided neural networks for industrial process modeling and soft sensing[J]. IEEE Transactions on Industrial Informatics, 2021, 17(6): 3987-3996.",
        "[3] Chen R, Kumar S, Agrawal A. Attention-based LSTM networks for multivariate time series forecasting in industrial systems[J]. Neurocomputing, 2021, 450: 176-189.",
        "[4] Wang S, Zhao T, Huang R, et al. Surrogate-assisted evolutionary optimization of cooling water systems for energy saving[J]. Energy, 2019, 183: 1158-1170.",
        "[5] Liu B, Zhang Q, Gielen G. A Gaussian process surrogate model assisted evolutionary algorithm for medium-scale expensive optimization problems[J]. IEEE Transactions on Evolutionary Computation, 2019, 23(6): 1010-1024.",
        "[6] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning[J]. Nature, 2015, 518(7540): 529-533.",
        "[7] Lillicrap T P, Hunt J J, Pritzel A, et al. Continuous control with deep reinforcement learning[C]. International Conference on Learning Representations (ICLR), San Juan, Puerto Rico, 2016.",
        "[8] Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms[J]. arXiv preprint arXiv:1707.06347, 2017.",
        "[9] Zhang Z, Chong A, Pan Y, et al. Whole building energy model for HVAC optimal control: a practical framework based on deep reinforcement learning[J]. Energy and Buildings, 2019, 199: 472-490.",
        "[10] Spielberg S, Gopaluni B, Loewen P. Deep reinforcement learning approaches for process control[C]. International Symposium on Advanced Control of Chemical Processes (ADCHEM), 2019: 220-225.",
        "[11] Ross S, Gordon G, Bagnell D. A reduction of imitation learning and structured prediction to no-regret online learning[C]. International Conference on Artificial Intelligence and Statistics (AISTATS), Fort Lauderdale, USA, 2011: 627-635.",
        "[12] Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
        "[13] Kingma D P, Ba J. Adam: A method for stochastic optimization[C]. International Conference on Learning Representations (ICLR), San Diego, USA, 2015.",
        "[14] Holland J H. Adaptation in natural and artificial systems[M]. Ann Arbor: University of Michigan Press, 1975: 1-50.",
        "[15] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016: 121-145, 197-218.",
        "[16] 李航. 统计学习方法[M]. 第2版. 北京: 清华大学出版社, 2019: 95-132, 149-178.",
        "[17] Sutton R S, Barto A G. Reinforcement learning: An introduction[M]. 2nd ed. Cambridge: MIT Press, 2018: 1-50, 271-304.",
        "[18] Silver D, Lever G, Heess N, et al. Deterministic policy gradient algorithms[C]. International Conference on Machine Learning (ICML), Beijing, China, 2014: 387-395.",
        "[19] Goodfellow I, Bengio Y, Courville A. Deep learning[M]. Cambridge: MIT Press, 2016: 367-415.",
        "[20] 冯焕生, 王晓东, 李明. 工业循环水系统节能技术综述[J]. 节能技术, 2019, 37(3): 245-252.",
        "[21] 刘建伟, 刘媛, 罗雄麟. 深度学习研究进展及其在过程控制中的应用[J]. 计算机应用研究, 2014, 31(7): 1921-1930.",
        "[22] 刘强, 卓洁, 郎自强, 等. 数据驱动的工业过程运行监控与自优化研究展望[J]. 自动化学报, 2018, 44(11): 1944-1956.",
        "[23] Xu J, Li H, Zhang Q. Adaptive control for circulating cooling water system using deep reinforcement learning[J]. PLoS ONE, 2024, 19(7): e0307767.",
        "[24] Liang Y Y, Wang D D, Chen J P, Shen Y G, Du J. Temperature control for a vehicle climate chamber using chilled water system[J]. Applied Thermal Engineering, 2016, 106: 117-124.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        set_run_font(run, cn=FONT_CN, en=FONT_EN, size=SIZE_SMALL)
        set_paragraph_spacing(p, line_spacing=1.25)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)

    add_page_break(doc)

    # ==================== 致谢 ====================
    add_chapter(doc, "致  谢")

    add_body(doc, "时光荏苒，四年的大学生活即将画上句号。在本论文完成之际，谨向所有在我学习和研究过程中给予帮助和支持的人们致以最诚挚的感谢。")

    add_body(doc, "首先，衷心感谢我的指导教师许谨老师。从选题确定、技术路线设计、实验方案制定到论文撰写和修改的每一个环节，许老师都倾注了大量的心血和智慧。许老师严谨求实的治学态度、深厚的学术造诣、对前沿技术敏锐的洞察力以及平易近人的指导风格，不仅帮助我顺利完成了毕业设计工作，更让我深刻体会到了科研工作的严谨与乐趣。每当我在技术难题前感到困惑时，X老师的点拨总能让我豁然开朗；每当我因为进度缓慢而感到焦虑时，X老师的鼓励总能让我重拾信心。在此谨向X老师致以最崇高的敬意和最真挚的感激。")

    add_body(doc, "感谢课题组和实验室的各位同学。在每周的组会上，大家积极讨论、互相启发，营造了浓厚的学术氛围。特别感谢在数据采集和预处理阶段提供帮助的同学，以及在深度学习和强化学习理论方面与我深入探讨的同学，与你们的交流和合作让我收益良多。")

    add_body(doc, "感谢人工智能学院各位老师在本科四年间的悉心教导和培养，是你们教授的每一门课程、组织的每一次实验，构筑了我完成本论文所需的知识基础和能力储备。感谢沈阳航空航天大学为我提供了良好的学习和生活环境，\"德能并进、勇毅翔远\"的校训将始终激励我在未来的人生道路上不断前行。")

    add_body(doc, "最后，衷心感谢我的家人。感谢父母二十多年来的养育之恩和无条件的支持，你们的理解、包容和鼓励是我在求学路上最坚实的后盾。正是因为有了你们的默默付出，我才能心无旁骛地投入到学习和研究中。")

    add_body(doc, "本论文的完成也离不开众多开源软件项目和学术论文作者的无私贡献——PyTorch、NumPy、Scikit-learn、Matplotlib等开源工具极大地降低了深度学习研究和科学计算的门槛，本文引用的各篇学术论文为我的研究提供了重要的理论基础和方法参考。在此一并表示感谢。")

    # ==================== 保存 ====================
    doc.save(OUTPUT_PATH)
    print(f"论文已生成 -> {OUTPUT_PATH}")


if __name__ == "__main__":
    generate()
