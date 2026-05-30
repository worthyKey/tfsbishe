from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r'D:\project\baogao\人工智能学院毕业设计（论文）任务书-基于数据驱动的循环水控制方法研究-滕凤硕.docx')

with open(r'D:\project\baogao\taskbook_full.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f'[P{i}] {p.text.strip()}\n')

    for ti, table in enumerate(doc.tables):
        f.write(f'\n=== Table {ti} ===\n')
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f'  Row{ri}: ' + ' | '.join(cells) + '\n')

print('Done - written to taskbook_full.txt')
